# -*- coding: utf-8 -*-
"""
make_scrap_docx.py — สร้าง Scrap_Module.docx (9-col standard, S1-S12 ครบ)
Format: เหมือน Maintain_Module.docx (Advisor Gen EXC-2)
Run: python -X utf8 docs/make_scrap_docx.py
"""
import sys, os
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color palette ──────────────────────────────────────────────────────────────
HDR_BG   = "1E3A5F"   # dark navy — header row
WHITE    = "FFFFFF"

CASE_BG = {
    "overview": "E8F0FE",  # blue-grey for S1-S12 overview
    "S1":  "FFF9C4",
    "S2":  "E8F5E9",
    "S3":  "E3F2FD",
    "S4":  "F3E5F5",
    "S5":  "FFF3E0",
    "S6":  "E1F5FE",
    "S7":  "FCE4EC",
    "S8":  "EFEBE9",
    "S9":  "F9FBE7",
    "S10": "E8EAF6",
    "S11": "FFEBEE",
    "S12": "FFF8E1",
    "reg": "F5F5F5",  # screen registry
}

HEADERS = [
    "รหัสจอ",
    "ชื่อจอ / หน้าที่",
    "มาจาก",
    "เงื่อนไข / เคส",
    "ไปต่อ",
    "แอพฯที่เห็น",
    "เคส",
    "หมายเหตุ",
    "ลิงก์ local",
]

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # replace existing shd if any
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    """Set approximate column widths (Cm list)."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def make_table(doc, rows_data: list[list[str]], bg_key: str) -> None:
    """
    rows_data[0] = header row (use HEADERS)
    rows_data[1:] = data rows
    bg_key = key in CASE_BG for data row background
    """
    n_cols = len(HEADERS)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, HEADERS)):
        cell.text = hdr
        set_cell_bg(cell, HDR_BG)
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(hdr)
        run.text = hdr
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)

    # Data rows
    data_bg = CASE_BG.get(bg_key, WHITE)
    for row_vals in rows_data:
        row = table.add_row()
        for i, (cell, val) in enumerate(zip(row.cells, row_vals)):
            cell.text = val
            set_cell_bg(cell, data_bg)
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(val)
            run.text = val
            run.font.size = Pt(8)

    # Column widths (A4 landscape ≈ 27.7 cm usable)
    # cols: [1.2, 3.5, 3.0, 3.5, 3.0, 2.8, 1.0, 3.2, 5.5] = ~26.7
    widths = [1.2, 3.5, 3.0, 3.5, 3.0, 2.8, 1.0, 3.2, 5.5]
    set_col_widths(table, widths)

    doc.add_paragraph("")  # spacing


def add_section_heading(doc, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.size = Pt(11 if level == 2 else 10)


# ── Data ───────────────────────────────────────────────────────────────────────

# §0 — S1-S12 Cases Overview (replaces old 2-col table)
OVERVIEW_ROWS = [
    # รหัสจอ | ชื่อจอ / หน้าที่ | มาจาก | เงื่อนไข / เคส | ไปต่อ | แอพฯที่เห็น | เคส | หมายเหตุ | ลิงก์ local
    ["S1", "ขายซาก (WeeeU → WeeeR)",
     "U-55 /scrap",
     "WeeeU ประกาศขาย (listingType=sell) · WeeeR ยื่น offer → WeeeU ยอมรับ → WeeeT รับซาก",
     "R-28b/c/d/e",
     "WeeeU · WeeeR · WeeeT",
     "S1", "Main flow — Escrow Gold Point lock",
     "http://localhost:3002/scrap"],
    ["S2", "ทิ้งซากฟรี (WeeeU → WeeeR ฟรี)",
     "U-55 /scrap",
     "WeeeU เลือก listingType=dispose isFree=true · WeeeR รับทำลาย/แยก",
     "R-28e /dispose",
     "WeeeU · WeeeR · WeeeT",
     "S2", "Dispose flow — ไม่มี Escrow",
     "http://localhost:3002/scrap"],
    ["S3", "WeeeR ซ่อมซากขาย (repair_and_sell)",
     "R-28 Job Detail",
     "WeeeR เลือก repair_and_sell → สร้าง RepairJob ย่อย → WeeeT ซ่อม → ขาย Marketplace",
     "R-28d → R-11 (Repair)",
     "WeeeR · WeeeT",
     "S3", "Cross-module: Repair sub-flow",
     "http://localhost:3001/scrap/jobs"],
    ["S4", "WeeeR เลือก dispose → E-Waste cert",
     "R-28 Job Detail",
     "WeeeR เลือก dispose → Admin ออก E-Waste cert → WeeeU รับใบรับรอง",
     "R-28e → A-11 → U-32",
     "WeeeR · Admin · WeeeU",
     "S4", "cert: EW-2026-xxxx",
     "http://localhost:3001/scrap/jobs"],
    ["S5", "ประกาศหมดอายุ → WeeeU ลงใหม่",
     "U-55 /scrap",
     "listing status=expired · ไม่มีร้านยื่น offer ภายในกำหนด",
     "U-29 /scrap/new (ลงใหม่)",
     "WeeeU",
     "S5", "MOCK: SCR-003 expired · ปุ่ม 'ลงใหม่'",
     "http://localhost:3002/scrap"],
    ["S6", "WeeeT รับซาก (Pickup confirm)",
     "T-22 /scrap",
     "WeeeT ได้รับ assignment pickup → ถึงที่รับ → ยืนยันรับ → status in_progress",
     "T-22 /scrap (กลับ list)",
     "WeeeT · WeeeU",
     "S6", "label 'ถึงจุดรับ' → 'ใบรับของ'",
     "http://localhost:3003/jobs"],
    ["S7", "WeeeR ถอนตัวหลังยืนยัน (canWithdraw)",
     "R-27 /scrap/jobs",
     "canWithdraw=true · WeeeR กด 'ถอนตัว' → Escrow คืน · status=cancelled",
     "R-27 /scrap/jobs",
     "WeeeR · WeeeU",
     "S7", "SJ005 resell_as_scrap/cancelled",
     "http://localhost:3001/scrap/jobs"],
    ["S8", "WeeeT พบซากไม่ตรงประกาศ (mismatch)",
     "T-04 /pickup",
     "WeeeT พบซากไม่ตรง → แจ้ง mismatch + เสนอราคาใหม่ → WeeeU ยินยอม/โต้แย้ง",
     "U-33 (ยินยอม/dispute)",
     "WeeeT · WeeeU · WeeeR",
     "S8", "mismatchReport: บันทึกเหตุผล + เสนอราคา",
     "http://localhost:3003/jobs"],
    ["S9", "WeeeT ถึงหน้างาน ไม่พบเจ้าของ (no-show)",
     "T-04 /pickup",
     "WeeeT รอ 20 นาที ไม่พบเจ้าของ → รายงาน no-show → WeeeU นัดใหม่/ยกเลิก",
     "U-33 (นัดใหม่/ยกเลิก)",
     "WeeeT · WeeeU",
     "S9", "ค่าเสียเที่ยวอาจถูกหักจาก Escrow",
     "http://localhost:3003/jobs"],
    ["S10", "WeeeU ยกเลิกหลัง in_progress",
     "U-55 /scrap",
     "listing status=in_progress · WeeeU กด 'ยกเลิก' → dialog เหตุผล → status cancelled",
     "U-55 /scrap",
     "WeeeU · WeeeR",
     "S10", "อาจมีค่าปรับ Escrow ตามเงื่อนไข",
     "http://localhost:3002/scrap"],
    ["S11", "WeeeU/WeeeR เปิด dispute (Escrow พิพาท)",
     "R-27 /scrap/jobs",
     "escrowStatus=locked · disputeReason ≠ null → Admin ตัดสินปล่อย Escrow",
     "A-09 /scrap/disputes",
     "WeeeR · WeeeU · Admin",
     "S11", "SJ006 badge 'พิพาท'",
     "http://localhost:3001/scrap/jobs"],
    ["S12", "Cross-module: Repair C4 → WeeeU ประกาศซาก",
     "U-07 /repair/c001/scrap-offer",
     "WeeeT วินิจฉัยซ่อมไม่คุ้ม → WeeeU เลือก 'ทิ้งซาก' → U-29 ?from_repair=REP-xxx",
     "U-29 → U-55 → R-70",
     "WeeeU · WeeeT · WeeeR",
     "S12", "Cross-module: Repair C4 flow",
     "http://localhost:3002/repair"],
]

# S1 — ขายซาก (WeeeU → WeeeR → WeeeT)
S1_ROWS = [
    ["U-29", "ประกาศซากใหม่", "U-55 /scrap",
     "WeeeU ต้องการขายซาก (listingType=sell)",
     "U-41 /scrap/new/success", "WeeeU (สร้าง)", "S1",
     "ตัวอย่าง: ตู้เย็น Samsung 2 ประตู · grade_B · ราคา 800 pts",
     "http://localhost:3002/scrap/new"],
    ["U-55", "ซากของฉัน (My Listings)", "U-41 success / U-01 dashboard",
     "หลัง publish listing ใหม่ / หน้ารายการ My Scrap",
     "U-33 /scrap/[id] (pending_offer)", "WeeeU (เจ้าของ)", "S1",
     "MOCK: SCR-001 pending_offer · 3 ข้อเสนอรอ",
     "http://localhost:3002/scrap"],
    ["R-70", "Scrap Hub (feed)", "sidebar WeeeR",
     "WeeeR ค้นหาซากที่น่าซื้อ",
     "R-78 /scrap/browse/[id]", "WeeeR (ผู้ซื้อ)", "S1",
     "MOCK_ITEMS SC001-SC006 · กรองตาม grade/ราคา",
     "http://localhost:3001/scrap"],
    ["R-72", "เลือกซื้อซาก (Browse)", "R-70 /scrap",
     "WeeeR เปิด browse list",
     "R-78 /scrap/browse/[id]", "WeeeR", "S1",
     "filter: grade_A/B/C + ราคา",
     "http://localhost:3001/scrap/browse"],
    ["R-78", "รายละเอียดซาก (Browse Detail)", "R-72 /scrap/browse",
     "WeeeR เลือกดูรายการ",
     "R-28 /scrap/jobs/[id] (ซื้อ)", "WeeeR (ผู้ซื้อ) · WeeeU (เจ้าของเห็น)", "S1",
     "'ซื้อซากนี้' → สร้าง ScrapJob · MOCK_ITEM SC002 Daikin",
     "http://localhost:3001/scrap/browse"],
    ["R-28", "งานซาก (Job Detail)", "R-78 browse / R-27 jobs",
     "งาน pending_decision ใหม่",
     "R-28b/c/d/e (ตามที่เลือก)", "WeeeR · WeeeU (เห็นสถานะ)", "S1",
     "เลือก 'ขายต่อซาก' → R-28b",
     "http://localhost:3001/scrap/jobs"],
    ["R-28b", "ตัดสินใจ: ขายต่อซาก", "R-28",
     "S1 เลือก resell_as_scrap",
     "R-27 /scrap/jobs (back)", "WeeeR", "S1",
     "ลง Marketplace ขายซากต่อ",
     "http://localhost:3001/scrap/jobs"],
    ["U-30", "ข้อเสนอรับซาก", "U-33 /scrap/[id]",
     "listing status=pending_offer · มี offer ≥1",
     "U-31 /scrap/[id]/confirm", "WeeeU (เจ้าของ) · WeeeR (ผู้ยื่น)", "S1",
     "MOCK_OFFERS: 3 ร้าน · countdown 24 ชม.",
     "http://localhost:3002/scrap/SCR-001/offers"],
    ["U-31", "ยืนยันเลือกข้อเสนอ", "U-30 /offers หรือ U-33 direct",
     "WeeeU กด 'เลือกข้อเสนอนี้'",
     "U-55 /scrap (status→accepted)", "WeeeU · WeeeR (รับแจ้ง) · WeeeT", "S1",
     "Escrow: WeeeR lock Gold Point ก่อน confirm",
     "http://localhost:3002/scrap/SCR-001/confirm"],
]

# S2 — ทิ้งซากฟรี
S2_ROWS = [
    ["U-29", "ประกาศซากใหม่", "U-55 /scrap",
     "WeeeU เลือก listingType=dispose isFree=true",
     "U-41 success", "WeeeU", "S2",
     "isFree=true · grade_C แนะนำ",
     "http://localhost:3002/scrap/new"],
    ["R-78", "รายละเอียดซาก", "R-72 /scrap/browse",
     "WeeeR เห็น badge 'ฟรี'",
     "R-28 /scrap/jobs/[id]", "WeeeR", "S2",
     "isFree=true · SC003 LG ตู้เย็น grade_C",
     "http://localhost:3001/scrap/browse"],
    ["R-28", "งานซาก", "R-78",
     "S2: dispose option",
     "R-28e /dispose", "WeeeR", "S2",
     "เลือก 'รีไซเคิล' ฟรี → R-28e",
     "http://localhost:3001/scrap/jobs"],
    ["U-30", "ข้อเสนอรับซาก", "U-33",
     "มี offer isFree ≥1",
     "U-31 confirm", "WeeeU", "S2",
     "MOCK offer type=free · ออก cert",
     "http://localhost:3002/scrap/SCR-003/offers"],
]

# S3 — ซ่อมขาย
S3_ROWS = [
    ["R-28", "งานซาก", "R-27 jobs",
     "S3: pending_decision",
     "R-28d /repair-and-sell", "WeeeR", "S3",
     "เลือก 'ซ่อมขาย' → R-28d",
     "http://localhost:3001/scrap/jobs"],
    ["R-28d", "ตัดสินใจ: ซ่อมขาย", "R-28",
     "S3 repair_and_sell → กรอก repair job form · เลือกช่าง WeeeT",
     "R-27 /scrap/jobs", "WeeeR · WeeeT (รับงานซ่อม)", "S3",
     "สร้าง repair_job ย่อย → WeeeT ซ่อม → ขาย Marketplace",
     "http://localhost:3001/scrap/jobs"],
]

# S4 — E-Waste cert
S4_ROWS = [
    ["R-28", "งานซาก", "R-27 jobs",
     "S4: pending_decision",
     "R-28e /dispose", "WeeeR", "S4",
     "เลือก 'รีไซเคิล' + E-Waste",
     "http://localhost:3001/scrap/jobs"],
    ["R-28e", "ตัดสินใจ: รีไซเคิล", "R-28",
     "S4 dispose → ขอ cert จาก Admin",
     "Admin A-11", "WeeeR · Admin", "S4",
     "WeeeR ส่ง request ออก E-Waste cert",
     "http://localhost:3001/scrap/jobs"],
    ["A-11", "ใบรับรอง E-Waste (Admin)", "Admin /scrap/jobs",
     "Admin เห็น dispose request",
     "A-11 /scrap/certificates/[id]", "Admin", "S4",
     "Admin พิมพ์ cert EW-2026-xxxx",
     "http://localhost:3000/scrap/certificates"],
    ["U-32", "ใบรับรอง E-Waste (WeeeU)", "U-33 /scrap/[id]",
     "cert ออกแล้ว · WeeeU เห็นปุ่มดู cert",
     "— (ดาวน์โหลด PDF)", "WeeeU (รับใบรับรอง) · Admin (ออก cert)", "S4",
     "MOCK_CERT: EW-2026-001234",
     "http://localhost:3002/scrap/SCR-003/certificate"],
]

# S5 — หมดอายุ ลงใหม่
S5_ROWS = [
    ["U-55", "ซากของฉัน", "dashboard / notification",
     "listing status=expired · ไม่มีร้านยื่น offer ภายในกำหนด",
     "U-29 /scrap/new (ลงใหม่)", "WeeeU", "S5",
     "MOCK: SCR-003 expired · ปุ่ม 'ลงใหม่'",
     "http://localhost:3002/scrap"],
    ["U-29", "ประกาศซากใหม่", "U-55 (กด 'ลงใหม่')",
     "S5 re-post — pre-fill ข้อมูลเดิม",
     "U-41 success", "WeeeU", "S5",
     "ระบบ reset expiry ใหม่",
     "http://localhost:3002/scrap/new"],
]

# S6 — WeeeT รับซาก
S6_ROWS = [
    ["T-04", "รับซาก (Pickup)", "T-22 /scrap หรือ T-01 /jobs",
     "WeeeT ได้รับ assignment pickup",
     "T-22 /scrap (กลับ list)", "WeeeT (ช่าง) · WeeeU (เห็น in_progress)", "S6",
     "label 'ถึงจุดรับ' → 'ใบรับของ'",
     "http://localhost:3003/jobs/J001/pickup"],
    ["U-33", "รายละเอียดซาก (WeeeU)", "push notification",
     "S6: status in_progress · ช่างกำลังเดินทาง",
     "U-55 /scrap (ปุ่มยกเลิก S10)", "WeeeU (เจ้าของ) · WeeeT (ช่าง)", "S6",
     "แสดง 'ช่างกำลังเดินทางมารับ'",
     "http://localhost:3002/scrap/SCR-002"],
]

# S7 — ถอนตัว
S7_ROWS = [
    ["R-27", "งานซากของฉัน", "R-70 /scrap",
     "MOCK SJ005: status=cancelled badge 'ถอนแล้ว'",
     "R-28 /scrap/jobs/[id]", "WeeeR", "S7",
     "canWithdraw=true ก่อนถอน · Escrow คืน",
     "http://localhost:3001/scrap/jobs"],
    ["R-28", "งานซาก (Job Detail)", "R-27 jobs",
     "S7: ปุ่ม 'ถอนตัว' (canWithdraw=true)",
     "R-27 /scrap/jobs", "WeeeR · WeeeU (รับ Escrow คืน)", "S7",
     "SJ005 resell_as_scrap/cancelled",
     "http://localhost:3001/scrap/jobs/SJ005"],
]

# S8 — mismatch
S8_ROWS = [
    ["T-10", "รายงานซากไม่ตรง (WeeeT)", "T-04 /pickup",
     "WeeeT พบซากไม่ตรงกับประกาศ",
     "T-22 /scrap (หลังส่งรายงาน)", "WeeeT (ช่าง) · WeeeU (รับแจ้งเตือน)", "S8",
     "mismatch: บันทึกเหตุผล + เสนอราคาใหม่",
     "http://localhost:3003/jobs/J001/mismatch"],
    ["U-33", "รายละเอียดซาก (WeeeU)", "push notification",
     "S8: mismatchReport ≠ null · banner แจ้งเตือน",
     "U-30 /offers (โต้แย้ง) หรือ ยินยอม", "WeeeU (ตัดสินใจ) · WeeeR (เห็น re-offer)", "S8",
     "แสดง mismatchReport: ช่าง+เหตุผล+รูป",
     "http://localhost:3002/scrap/SCR-002"],
    ["R-28", "งานซาก (WeeeR)", "R-27 jobs",
     "S8: status=pending_decision (back) after mismatch",
     "R-28c /resell-parts", "WeeeR", "S8",
     "SJ002: pending_decision หลัง mismatch · badge ⚠️",
     "http://localhost:3001/scrap/jobs/SJ002"],
]

# S9 — no-show
S9_ROWS = [
    ["T-04", "รับซาก — No Show", "T-01 /jobs",
     "WeeeT รอ 20 นาที ไม่พบเจ้าของ → รายงาน no-show",
     "T-22 /scrap", "WeeeT (รายงาน) · WeeeU (รับแจ้งเตือน)", "S9",
     "label 'ถึงจุดรับ' ไม่พบ → report button",
     "http://localhost:3003/jobs/J001/pickup"],
    ["U-33", "รายละเอียดซาก (WeeeU)", "push notification",
     "S9: noShowEvent ≠ null",
     "ปุ่ม 'นัดใหม่' / 'ยกเลิก'", "WeeeU (ตัดสินใจ) · WeeeT (รอ)", "S9",
     "แสดง 'ช่างมาถึงแล้ว แต่ไม่พบคุณ' · ค่าเสียเที่ยวจาก Escrow",
     "http://localhost:3002/scrap/SCR-002"],
]

# S10 — ยกเลิก in_progress
S10_ROWS = [
    ["U-55", "ซากของฉัน", "dashboard",
     "listing status=in_progress · ปุ่ม 'ยกเลิก'",
     "U-33 /scrap/[id]?action=cancel", "WeeeU", "S10",
     "MOCK SCR-002 in_progress · link action=cancel",
     "http://localhost:3002/scrap"],
    ["U-33", "รายละเอียดซาก + Cancel Dialog", "U-55 (link action=cancel)",
     "actionParam=cancel → dialog เปิดอัตโนมัติ",
     "U-55 /scrap (status→cancelled)", "WeeeU (ยืนยัน) · WeeeR (รับแจ้ง)", "S10",
     "dialog: กรอกเหตุผล บังคับ · อาจมีค่าปรับ",
     "http://localhost:3002/scrap/SCR-002"],
]

# S11 — dispute
S11_ROWS = [
    ["R-27", "งานซากของฉัน", "R-70 /scrap",
     "SJ006 badge 'พิพาท' · escrowStatus=locked+dispute",
     "R-28 /scrap/jobs/SJ006", "WeeeR", "S11",
     "disputeReason: WeeeU ส่งซากผิดชิ้น",
     "http://localhost:3001/scrap/jobs"],
    ["R-28", "งานซาก (Dispute)", "R-27",
     "S11: escrowStatus=locked · disputeReason ≠ null",
     "A-09 /scrap/disputes (Admin)", "WeeeR · Admin", "S11",
     "แสดง 'Escrow พิพาท' · รอ Admin ตัดสิน",
     "http://localhost:3001/scrap/jobs/SJ006"],
    ["A-09", "Dispute List (Admin)", "Admin sidebar",
     "Admin เห็น dispute ที่ต้องจัดการ",
     "A-10 /scrap/disputes/[id]", "Admin", "S11",
     "รายการ disputes รออนุมัติ",
     "http://localhost:3000/scrap/disputes"],
    ["A-10", "Dispute Detail (Admin)", "A-09 /scrap/disputes",
     "Admin ตรวจสอบหลักฐานทั้งสองฝ่าย",
     "A-09 (กลับ list)", "Admin · WeeeR · WeeeU", "S11",
     "ตัดสิน: ปล่อย Escrow ให้ใคร",
     "http://localhost:3000/scrap/disputes"],
]

# S12 — Cross-module Repair C4
S12_ROWS = [
    ["U-07", "Repair C4 Scrap Offer", "U-06 /repair/[id]/progress",
     "WeeeT วินิจฉัย: ซ่อมไม่คุ้ม → WeeeU เลือก 'ทิ้งซาก' / 'ขายซาก'",
     "U-29 /scrap/new?from_repair=REP-xxx", "WeeeU (ตัดสินใจ) · WeeeT (T-06)", "S12",
     "C4 flow: WeeeT ใช้ T-06 ส่งข้อมูลราคาประเมิน",
     "http://localhost:3002/repair"],
    ["T-06", "Repair C4 Scrap Offer (WeeeT)", "T-02 /diagnose",
     "S12: WeeeT กรอก assessed price + weight",
     "T-11 /jobs/[id]", "WeeeT", "S12",
     "estimated price + weight → ส่งให้ WeeeU",
     "http://localhost:3003/jobs/c001/scrap-offer"],
    ["U-29", "ประกาศซากใหม่ (S12 pre-fill)", "U-07 /repair/c001/scrap-offer",
     "?from_repair=REP-xxx → banner + pre-fill ข้อมูลจาก Repair",
     "U-41 success", "WeeeU", "S12",
     "banner 'งานซ่อม → แนะนำทิ้งซาก' · ราคาประเมินจาก WeeeT",
     "http://localhost:3002/scrap/new?from_repair=REP-0042"],
    ["U-55", "ซากของฉัน", "U-41 success",
     "listing มี sourceRepairJobId",
     "U-33 /scrap/[id]", "WeeeU", "S12",
     "MOCK SCR-002 badge 'งานซ่อม #REP-0042'",
     "http://localhost:3002/scrap"],
    ["U-33", "รายละเอียดซาก (S12)", "U-55 /scrap",
     "sourceRepairJobId ≠ null",
     "U-30 offers / U-31 confirm", "WeeeU · WeeeR", "S12",
     "header badge 'งานซ่อม #REP-0042'",
     "http://localhost:3002/scrap/SCR-002"],
    ["R-70", "Scrap Hub Feed", "sidebar",
     "S12 listing มี badge 'จาก Repair'",
     "R-78 /scrap/browse/[id]", "WeeeR", "S12",
     "MOCK SC004 HP Notebook fromRepair",
     "http://localhost:3001/scrap"],
    ["T-22", "Scrap Jobs (WeeeT)", "T-01 /jobs",
     "WeeeT เห็น scrap pickup job จาก S12",
     "T-23 /scrap/[id]", "WeeeT", "S12",
     "T-22: WeeeT scrap list",
     "http://localhost:3003/scrap"],
    ["T-23", "Scrap Job Detail (WeeeT)", "T-22 /scrap",
     "S12 job detail · badge Repair#REP-0042",
     "T-04 /jobs/[id]/pickup", "WeeeT", "S12",
     "T-23: รายละเอียด scrap job",
     "http://localhost:3003/scrap/J001"],
]

# Screen Registry (9-col: รหัสจอ / ชื่อจอ / แอพ / Route / เคส / แอพฯที่เห็น / ไปต่อ / หมายเหตุ / ลิงก์ local)
# Map: [รหัสจอ, ชื่อจอ / หน้าที่, มาจาก (แอพ), เงื่อนไข / เคส (Code), ไปต่อ, แอพฯที่เห็น, เคส, หมายเหตุ (route+note), ลิงก์ local]
REGISTRY_ROWS = [
    ["U-29", "ประกาศซากใหม่", "WeeeU", "SCRAP-CREATE", "U-41 success", "WeeeU", "S1-S5,S12", "/scrap/new · S1-S5 entry · S12 ?from_repair", "http://localhost:3002/scrap/new"],
    ["U-30", "ข้อเสนอรับซาก", "WeeeU", "SCRAP-S1-OFFERS", "U-31 confirm", "WeeeU · WeeeR", "S1,S2", "/scrap/[id]/offers · S1/S2 offer list", "http://localhost:3002/scrap/SCR-001/offers"],
    ["U-31", "ยืนยันเลือกข้อเสนอ", "WeeeU", "SCRAP-S1-CONFIRM", "U-55 /scrap", "WeeeU · WeeeR · WeeeT", "S1,S2", "/scrap/[id]/confirm · ยืนยัน S1/S2", "http://localhost:3002/scrap/SCR-001/confirm"],
    ["U-32", "ใบรับรอง E-Waste (WeeeU)", "WeeeU", "SCRAP-S4-CERT", "— (download)", "WeeeU · Admin", "S4", "/scrap/[id]/certificate · S4 E-Waste cert", "http://localhost:3002/scrap/SCR-003/certificate"],
    ["U-33", "รายละเอียดซาก (WeeeU)", "WeeeU", "SCRAP-DETAIL", "U-30 offers / U-31 confirm", "WeeeU · WeeeR · WeeeT", "S1-S5,S8-S10,S12", "/scrap/[id] · S1-S5·S8·S9·S10·S12", "http://localhost:3002/scrap/SCR-001"],
    ["U-41", "ประกาศซากสำเร็จ", "WeeeU", "SCRAP-CREATE-SUCCESS", "U-55 /scrap", "WeeeU", "S1,S2,S12", "/scrap/new/success · หลัง submit", "http://localhost:3002/scrap/new/success"],
    ["U-55", "ซากของฉัน (My Listings)", "WeeeU", "SCRAP-HOME", "U-33 /scrap/[id]", "WeeeU", "S1-S7,S10,S12", "/scrap · My listings list", "http://localhost:3002/scrap"],
    ["R-24", "Scrap Announce List", "WeeeR", "SCRAP-ANNOUNCE-LIST", "R-72 /scrap/browse", "WeeeR", "S1,S2", "/scrap/announcements → redirect R-72", "http://localhost:3001/scrap/announcements"],
    ["R-25", "ยื่น Offer (Scrap)", "WeeeR", "SCRAP-BID", "U-30 offers (WeeeU)", "WeeeR", "S1,S2", "/scrap/announcements/[id]/offer · ยื่น offer", "http://localhost:3001/scrap/announcements"],
    ["R-26", "รายละเอียดประกาศซาก", "WeeeR", "SCRAP-ANNOUNCE-DETAIL", "R-78 /scrap/browse/[id]", "WeeeR", "S1,S2", "/scrap/announcements/[id]", "http://localhost:3001/scrap/announcements"],
    ["R-27", "งานซากของฉัน (Jobs List)", "WeeeR", "SCRAP-JOBS", "R-28 /scrap/jobs/[id]", "WeeeR", "S1-S12", "/scrap/jobs · งานซาก S1-S12", "http://localhost:3001/scrap/jobs"],
    ["R-28", "งานซาก (Job Detail)", "WeeeR", "SCRAP-JOB-DETAIL", "R-28b/c/d/e", "WeeeR · WeeeU", "S1-S4,S7,S8,S11,S12", "/scrap/jobs/[id] · ตัดสินใจ+badge", "http://localhost:3001/scrap/jobs/SJ001"],
    ["R-28b", "ตัดสินใจ: ขายต่อซาก", "WeeeR", "SCRAP-S1-DECISION", "R-27 /scrap/jobs", "WeeeR", "S1", "/scrap/jobs/[id]/resell-as-scrap", "http://localhost:3001/scrap/jobs/SJ001/resell-as-scrap"],
    ["R-28c", "ตัดสินใจ: แยกอะไหล่", "WeeeR", "SCRAP-S2-DECISION", "R-27 /scrap/jobs", "WeeeR", "S2", "/scrap/jobs/[id]/resell-parts", "http://localhost:3001/scrap/jobs/SJ001/resell-parts"],
    ["R-28d", "ตัดสินใจ: ซ่อมขาย", "WeeeR", "SCRAP-S3-DECISION", "R-11 /repair/jobs/[id]", "WeeeR · WeeeT", "S3", "/scrap/jobs/[id]/repair-and-sell", "http://localhost:3001/scrap/jobs/SJ001/repair-and-sell"],
    ["R-28e", "ตัดสินใจ: รีไซเคิล/E-Waste", "WeeeR", "SCRAP-S4-DECISION", "A-11 /scrap/certificates", "WeeeR · Admin", "S4", "/scrap/jobs/[id]/dispose", "http://localhost:3001/scrap/jobs/SJ001/dispose"],
    ["R-70", "Scrap Hub Feed (Public)", "WeeeR", "SCRAP-HUB", "R-78 /scrap/browse/[id]", "WeeeR", "S1,S2,S12", "/scrap · Public feed", "http://localhost:3001/scrap"],
    ["R-71", "Scrap Item Detail (Public)", "WeeeR", "SCRAP-ITEM-DETAIL", "R-28 (ซื้อ)", "WeeeR", "S1,S2", "/scrap/[id] · Public detail", "http://localhost:3001/scrap/SCR-001"],
    ["R-72", "เลือกซื้อซาก (Browse+Filter)", "WeeeR", "SCRAP-BROWSE", "R-78 /scrap/browse/[id]", "WeeeR", "S1,S2", "/scrap/browse · grade/ราคา filter", "http://localhost:3001/scrap/browse"],
    ["R-78", "รายละเอียด + ปุ่มซื้อ", "WeeeR", "SCRAP-BROWSE-DETAIL", "R-28 /scrap/jobs/[id]", "WeeeR · WeeeU", "S1,S2", "/scrap/browse/[id] · ซื้อซากนี้", "http://localhost:3001/scrap/browse/SCR-001"],
    ["T-04", "รับซาก (Pickup / No Show)", "WeeeT", "SCRAP-S6-PICKUP", "T-22 /scrap", "WeeeT · WeeeU", "S6,S9", "/jobs/[id]/pickup · S6·S9", "http://localhost:3003/jobs/J001/pickup"],
    ["T-06", "Repair C4 Scrap Offer (WeeeT)", "WeeeT", "REPAIR-C4-SCRAP", "T-11 /jobs/[id]", "WeeeT · WeeeU", "S12", "/jobs/[id]/scrap-offer · S12 C4 assessed price", "http://localhost:3003/jobs/c001/scrap-offer"],
    ["T-10", "รายงานซากไม่ตรง (Mismatch)", "WeeeT", "SCRAP-S8-MISMATCH", "T-22 /scrap", "WeeeT · WeeeU · WeeeR", "S8", "/jobs/[id]/mismatch · S8", "http://localhost:3003/jobs/J001/mismatch"],
    ["T-22", "Scrap Jobs List (WeeeT)", "WeeeT", "SCRAP-HOME (T)", "T-23 /scrap/[id]", "WeeeT", "S6,S9,S12", "/scrap · WeeeT scrap list", "http://localhost:3003/scrap"],
    ["T-23", "Scrap Job Detail (WeeeT)", "WeeeT", "SCRAP-DETAIL (T)", "T-04 /pickup", "WeeeT", "S6,S9,S12", "/scrap/[id] · WeeeT job detail", "http://localhost:3003/scrap/J001"],
    ["A-08", "Scrap Jobs Overview (Admin)", "Admin", "SCRAP-JOBS-ADMIN", "A-09 /scrap/disputes", "Admin", "S1-S12", "/scrap/jobs · overview ทุก jobs", "http://localhost:3000/scrap/jobs"],
    ["A-09", "Dispute List (Admin)", "Admin", "SCRAP-DISPUTES", "A-10 /scrap/disputes/[id]", "Admin", "S11", "/scrap/disputes · S11 dispute list", "http://localhost:3000/scrap/disputes"],
    ["A-10", "Dispute Detail (Admin)", "Admin", "SCRAP-DISPUTE-DETAIL", "A-09 /scrap/disputes", "Admin · WeeeR · WeeeU", "S11", "/scrap/disputes/[id] · S11 ตัดสิน", "http://localhost:3000/scrap/disputes/D001"],
    ["A-11", "E-Waste Certificates (Admin)", "Admin", "SCRAP-CERTS", "A-11 /scrap/certificates/[id]", "Admin · WeeeU", "S4", "/scrap/certificates · S4 cert mgmt", "http://localhost:3000/scrap/certificates"],
]


# ── Build Document ──────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page layout: A4 landscape
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Title
    title = doc.add_heading("Scrap Module — Flow & Screen Reference", 0)
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("9-column standard format · S1-S12 ครบ · local links per app")

    # §0 S1-S12 Overview
    add_section_heading(doc, "§0  สรุป S1–S12 Cases (Overview)", 2)
    make_table(doc, OVERVIEW_ROWS, "overview")

    # §1–§12 per-case tables
    CASES = [
        ("S1", "ขายซาก — WeeeU ประกาศ → WeeeR ยื่น offer → WeeeT รับซาก", S1_ROWS),
        ("S2", "ทิ้งซากฟรี — WeeeU dispose → WeeeR รับไปทำลาย/แยก (isFree=true)", S2_ROWS),
        ("S3", "ซ่อมขาย — WeeeR เลือก repair_and_sell → ส่ง WeeeT ซ่อม → ขาย Marketplace", S3_ROWS),
        ("S4", "E-Waste cert — WeeeR เลือก dispose → Admin ออก cert → WeeeU รับ cert", S4_ROWS),
        ("S5", "ประกาศหมดอายุ — ไม่มีร้านสนใจ → WeeeU กด 'ลงใหม่'", S5_ROWS),
        ("S6", "WeeeT รับซาก (Pickup) — ถึงที่รับ → ยืนยันรับ → in_progress", S6_ROWS),
        ("S7", "WeeeR ถอนตัว — canWithdraw=true → Escrow คืน", S7_ROWS),
        ("S8", "Mismatch — WeeeT พบซากไม่ตรงประกาศ → WeeeU ยินยอม/โต้แย้ง", S8_ROWS),
        ("S9", "No-show — WeeeT ถึงหน้างาน ไม่พบเจ้าของ → นัดใหม่/ยกเลิก", S9_ROWS),
        ("S10", "WeeeU ยกเลิก — ระหว่าง in_progress → อาจมีค่าปรับ", S10_ROWS),
        ("S11", "Dispute — Escrow พิพาท → Admin ตัดสิน", S11_ROWS),
        ("S12", "Cross-module Repair C4 — WeeeT ซ่อมไม่คุ้ม → WeeeU ประกาศซาก", S12_ROWS),
    ]

    for case_key, case_desc, rows in CASES:
        add_section_heading(doc, f"§{case_key[1:]}  {case_key} — {case_desc}", 2)
        make_table(doc, rows, case_key)

    # Screen Registry
    add_section_heading(doc, "§13  Screen Registry (ทุก Screen ID ในโมดูล Scrap)", 2)
    make_table(doc, REGISTRY_ROWS, "reg")

    # Save
    out = os.path.join(os.path.dirname(__file__), "Scrap_Module.docx")
    doc.save(out)
    print(f"Saved: {out}")

    # Verify
    d2 = Document(out)
    total_rows = sum(len(t.rows) for t in d2.tables)
    cols = [len(t.columns) for t in d2.tables]
    print(f"rows {total_rows} cols {cols}")
    bad = [c for c in cols if c != 9]
    if bad:
        print(f"ERROR: tables with wrong cols: {bad}")
        sys.exit(1)
    else:
        print("✅ All tables = 9 columns")


if __name__ == "__main__":
    build()
