# -*- coding: utf-8 -*-
"""
make_parts_weeet_docx.py — สร้าง Parts_WeeeT.docx
Parts Request/Order Flow สำหรับช่าง WeeeT (Buyer Role)
Format: 9-col standard (เหมือน Scrap_Module.docx)
Run: python -X utf8 docs/make_parts_weeet_docx.py
"""
import sys, os
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Color palette ──────────────────────────────────────────────────────────────
HDR_BG  = "1A4D80"   # WeeeT blue — header row
WHITE   = "FFFFFF"

CASE_BG = {
    "overview": "E3F2FD",  # light blue — overview
    "B1": "E8F5E9",
    "B2": "FFF9C4",
    "B3": "FCE4EC",
    "B4": "F3E5F5",
    "B5": "FFF3E0",
    "reg": "F5F5F5",
}

HEADERS = [
    "รหัสจอ",
    "ชื่อจอ / หน้าที่",
    "มาจาก",
    "เงื่อนไข / เคส",
    "ไปต่อ",
    "แอพฯที่เห็น",
    "เคส",
    "หมายเหตุ",
    "ลิงก์ local",
]

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def make_table(doc, rows_data: list[list[str]], bg_key: str) -> None:
    n_cols = len(HEADERS)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"

    hdr_row = table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, HEADERS)):
        cell.text = hdr
        set_cell_bg(cell, HDR_BG)
        para = cell.paragraphs[0]
        run = para.runs[0] if para.runs else para.add_run(hdr)
        run.text = hdr
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)

    data_bg = CASE_BG.get(bg_key, WHITE)
    for row_vals in rows_data:
        row = table.add_row()
        for i, (cell, val) in enumerate(zip(row.cells, row_vals)):
            cell.text = val
            set_cell_bg(cell, data_bg)
            para = cell.paragraphs[0]
            run = para.runs[0] if para.runs else para.add_run(val)
            run.text = val
            run.font.size = Pt(8)

    widths = [1.2, 3.5, 3.0, 3.5, 3.0, 2.8, 1.0, 3.2, 5.5]
    set_col_widths(table, widths)
    doc.add_paragraph("")


def add_section_heading(doc, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.size = Pt(11 if level == 2 else 10)


# ── Data ───────────────────────────────────────────────────────────────────────

# §0 — Overview: B1-B5 Cases
OVERVIEW_ROWS = [
    # [รหัสจอ, ชื่อจอ/หน้าที่, มาจาก, เงื่อนไข/เคส, ไปต่อ, แอพฯที่เห็น, เคส, หมายเหตุ, ลิงก์ local]
    ["T-24", "Parts Hub (WeeeT)",
     "BottomNav / T-18 Dashboard",
     "WeeeT เปิดเมนูหลักอะไหล่",
     "T-26 /catalog · T-30 /orders · T-32 /requests",
     "WeeeT",
     "B1-B5", "entry point ทุก flow", "http://localhost:3003/parts"],
    ["T-26", "Parts Catalog",
     "T-24 /parts",
     "WeeeT เรียกดูอะไหล่ใน B2B marketplace",
     "T-27 /catalog/[id] · T-28 /cart",
     "WeeeT · WeeeR (เห็นยอดขาย)",
     "B1", "browse + search + filter", "http://localhost:3003/parts/catalog"],
    ["T-27", "Catalog Item Detail",
     "T-26 /catalog",
     "WeeeT ดูรายละเอียดอะไหล่ + tier pricing",
     "T-28 /cart (เพิ่มตะกร้า)",
     "WeeeT · WeeeR",
     "B1", "conditionScore / tierPricing", "http://localhost:3003/parts/catalog/PART-001"],
    ["T-28", "Shopping Cart",
     "T-27 /catalog/[id]",
     "WeeeT ตรวจสอบรายการ + จำนวน ก่อนชำระ",
     "T-29 /checkout",
     "WeeeT",
     "B1", "qty adjust / remove item", "http://localhost:3003/parts/cart"],
    ["T-29", "Checkout",
     "T-28 /cart",
     "WeeeT ยืนยันสั่งซื้อ · ชำระ 100 Gold/order",
     "T-31 /orders/[id] (order created)",
     "WeeeT · Admin (audit) · WeeeR (รับออเดอร์)",
     "B1", "fee = 100 Gold · idempotencyKey", "http://localhost:3003/parts/checkout"],
    ["T-30", "My Orders List",
     "T-24 /parts · BottomNav",
     "WeeeT ดูรายการออเดอร์ทั้งหมด",
     "T-31 /orders/[id]",
     "WeeeT",
     "B1,B3,B4", "filter by status", "http://localhost:3003/parts/orders"],
    ["T-31", "Order Detail + Audit Trail",
     "T-30 /orders",
     "WeeeT ดูรายละเอียด + timeline · ปิด/พิพาท/ให้คะแนน",
     "T-30 /orders (ปิดงาน) · T-30 (dispute)",
     "WeeeT · WeeeR · Admin",
     "B1,B3,B4", "closeOrder / disputeOrder / rateOrder", "http://localhost:3003/parts/orders/ORD-001"],
    ["T-32", "Parts Requests",
     "T-24 /parts",
     "WeeeT ส่งคำขออะไหล่ที่หาไม่เจอใน catalog",
     "T-30 /orders (เมื่อ seller fulfill)",
     "WeeeT · WeeeR (seller เห็น request)",
     "B2", "partRequest: partName + description + jobId", "http://localhost:3003/parts/requests"],
    ["T-36", "Parts for Job",
     "T-11 /jobs/[id]",
     "WeeeT เปิดดูอะไหล่สำหรับงานซ่อมนี้โดยเฉพาะ",
     "T-24 /parts (เพิ่มตะกร้า)",
     "WeeeT",
     "B5", "context: jobId filter", "http://localhost:3003/jobs/J001/parts"],
    ["T-25", "Part Detail (Direct Link)",
     "T-32 /requests (deep link)",
     "WeeeT เปิดอะไหล่ชิ้นเดียวตรง ๆ",
     "T-28 /cart",
     "WeeeT",
     "B1,B2", "direct part link from notification/request", "http://localhost:3003/parts/PART-001"],
]

# B1 — Browse → Cart → Checkout → Order (main flow)
B1_ROWS = [
    ["T-26", "Parts Catalog", "T-24 /parts",
     "WeeeT ค้นหาอะไหล่ใน B2B marketplace",
     "T-27 /catalog/[id]", "WeeeT", "B1",
     "ค้นหาด้วย partName / manufacturer / conditionScore filter",
     "http://localhost:3003/parts/catalog"],
    ["T-27", "Catalog Item Detail", "T-26 /catalog",
     "เปิดรายละเอียดอะไหล่ · ดู tier pricing / warranty",
     "T-28 /cart (กด เพิ่มตะกร้า)", "WeeeT · WeeeR (seller)", "B1",
     "MOCK: PART-001 คอมเพรสเซอร์ Daikin · 1,800 Gold · tier ≥3 ชิ้น -5%",
     "http://localhost:3003/parts/catalog/PART-001"],
    ["T-28", "Shopping Cart", "T-27 (เพิ่ม) หรือ T-26 (ตะกร้า chip)",
     "ตรวจสอบรายการ · ปรับจำนวน · ลบ",
     "T-29 /checkout", "WeeeT", "B1",
     "qty ≥ qtyAvailable → แจ้งเตือน · localStorage cart",
     "http://localhost:3003/parts/cart"],
    ["T-29", "Checkout", "T-28 /cart",
     "ยืนยันสั่งซื้อ · สร้าง order · หัก 100 Gold ค่าบริการ",
     "T-31 /orders/[id] (status=pending)", "WeeeT · WeeeR (รับออเดอร์) · Admin", "B1",
     "POST /api/v1/parts/orders/ · idempotencyKey = uuid",
     "http://localhost:3003/parts/checkout"],
    ["T-30", "My Orders List", "T-29 (หลัง checkout) · T-24 chip",
     "รายการออเดอร์ · filter: pending/confirmed/shipped/delivered/closed",
     "T-31 /orders/[id]", "WeeeT", "B1",
     "MOCK: ORD-001 pending · ORD-002 shipped",
     "http://localhost:3003/parts/orders"],
    ["T-31", "Order Detail", "T-30 /orders",
     "ดูรายละเอียด + audit trail · ปุ่ม ยืนยันรับ (status=delivered)",
     "T-30 /orders (กลับ list)", "WeeeT · WeeeR · Admin", "B1",
     "PATCH /parts/orders/:id/close/ · หลัง close → prompt rate",
     "http://localhost:3003/parts/orders/ORD-001"],
]

# B2 — Parts Request flow
B2_ROWS = [
    ["T-32", "Parts Requests", "T-24 /parts",
     "WeeeT ส่งคำขออะไหล่ที่หาไม่เจอใน catalog",
     "T-31 /orders/[id] (เมื่อ WeeeR fulfill)", "WeeeT · WeeeR", "B2",
     "partRequest: partName + description + qty + jobId (optional)",
     "http://localhost:3003/parts/requests"],
    ["T-36", "Parts for Job", "T-11 /jobs/[id]",
     "WeeeT เปิดระหว่างงานซ่อม · filter catalog ตาม jobId",
     "T-28 /cart (เพิ่มจาก job context)", "WeeeT", "B2,B5",
     "jobId parameter ส่งไปกับ order ใน createOrder()",
     "http://localhost:3003/jobs/J001/parts"],
    ["T-27", "Catalog Item Detail (from Request)", "notification / T-32 deep link",
     "WeeeR fulfill request → WeeeT รับ notification → เปิด part detail",
     "T-28 /cart", "WeeeT", "B2",
     "deep link: /parts/catalog/[id]?from_request=REQ-001",
     "http://localhost:3003/parts/catalog/PART-007"],
]

# B3 — Dispute Order
B3_ROWS = [
    ["T-31", "Order Detail — Dispute", "T-30 /orders",
     "WeeeT พบปัญหา (ของไม่ตรง/ชำรุด) · กด แจ้งข้อพิพาท",
     "T-30 /orders (status=disputed)", "WeeeT · WeeeR · Admin", "B3",
     "POST /parts/orders/:id/dispute/ · reason บังคับ",
     "http://localhost:3003/parts/orders/ORD-001"],
]

# B4 — Close & Rate
B4_ROWS = [
    ["T-31", "Order Detail — Rate", "T-30 /orders (status=closed)",
     "WeeeT ปิดงานแล้ว · prompt ให้คะแนน seller",
     "T-30 /orders", "WeeeT · WeeeR (เห็นคะแนน)", "B4",
     "POST /parts/orders/:id/rate/ · score 1-5 + comment optional",
     "http://localhost:3003/parts/orders/ORD-001"],
]

# B5 — Parts in Job Context (T-36)
B5_ROWS = [
    ["T-36", "Parts for Job", "T-11 /jobs/[id] (ปุ่ม อะไหล่)",
     "WeeeT ซ่อมแล้วพบต้องใช้อะไหล่เพิ่ม · เปิด T-36",
     "T-26 /catalog (filter jobId) · T-28 /cart", "WeeeT", "B5",
     "jobId ส่งใน createOrder() เพื่อ audit trail งานซ่อม",
     "http://localhost:3003/jobs/J001/parts"],
    ["T-29", "Checkout (Job context)", "T-28 /cart",
     "WeeeT checkout อะไหล่สำหรับงาน · serviceId = jobId",
     "T-31 /orders/[id]", "WeeeT · WeeeR · Admin", "B5",
     "createOrder({ partId, qty, serviceId: jobId }) · fee 100 Gold",
     "http://localhost:3003/parts/checkout"],
]

# Screen Registry
REGISTRY_ROWS = [
    ["T-24", "Parts Hub", "WeeeT", "PARTS-HUB", "T-26/T-30/T-32", "WeeeT", "B1-B5", "/parts · entry point", "http://localhost:3003/parts"],
    ["T-25", "Part Detail (Direct)", "WeeeT", "PARTS-ITEM-DIRECT", "T-28 /cart", "WeeeT", "B1,B2", "/parts/[id] · direct link", "http://localhost:3003/parts/PART-001"],
    ["T-26", "Parts Catalog", "WeeeT", "PARTS-CATALOG", "T-27 /catalog/[id]", "WeeeT · WeeeR", "B1,B2", "/parts/catalog · browse B2B", "http://localhost:3003/parts/catalog"],
    ["T-27", "Catalog Item Detail", "WeeeT", "PARTS-CATALOG-DETAIL", "T-28 /cart", "WeeeT · WeeeR", "B1,B2", "/parts/catalog/[id] · tier pricing", "http://localhost:3003/parts/catalog/PART-001"],
    ["T-28", "Shopping Cart", "WeeeT", "PARTS-CART", "T-29 /checkout", "WeeeT", "B1,B5", "/parts/cart · qty/remove", "http://localhost:3003/parts/cart"],
    ["T-29", "Checkout", "WeeeT", "PARTS-CHECKOUT", "T-31 /orders/[id]", "WeeeT · WeeeR · Admin", "B1,B5", "/parts/checkout · 100 Gold fee · idempotencyKey", "http://localhost:3003/parts/checkout"],
    ["T-30", "My Orders List", "WeeeT", "PARTS-ORDERS", "T-31 /orders/[id]", "WeeeT", "B1,B3,B4", "/parts/orders · status filter", "http://localhost:3003/parts/orders"],
    ["T-31", "Order Detail", "WeeeT", "PARTS-ORDER-DETAIL", "T-30 /orders", "WeeeT · WeeeR · Admin", "B1,B3,B4", "/parts/orders/[id] · close/dispute/rate", "http://localhost:3003/parts/orders/ORD-001"],
    ["T-32", "Parts Requests", "WeeeT", "PARTS-REQUESTS", "T-31 (fulfilled)", "WeeeT · WeeeR", "B2", "/parts/requests · ส่งคำขออะไหล่", "http://localhost:3003/parts/requests"],
    ["T-36", "Parts for Job", "WeeeT", "PARTS-JOB-CONTEXT", "T-26/T-28", "WeeeT", "B2,B5", "/jobs/[id]/parts · jobId context", "http://localhost:3003/jobs/J001/parts"],
]


# ── Build Document ──────────────────────────────────────────────────────────────

def build():
    doc = Document()

    from docx.enum.section import WD_ORIENT
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    title = doc.add_heading("Parts WeeeT — Flow & Screen Reference (Buyer Role)", 0)
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph("9-column standard format · B1-B5 cases · WeeeT ช่าง = Buyer role · ห้าม seller actions")

    # §0 Overview
    add_section_heading(doc, "§0  สรุป B1–B5 Cases + Screen Map (Overview)", 2)
    make_table(doc, OVERVIEW_ROWS, "overview")

    # §1–§5 per-case
    CASES = [
        ("B1", "Browse Catalog → Cart → Checkout → Order Tracking (main flow)", B1_ROWS),
        ("B2", "Parts Request — ช่างส่งคำขออะไหล่ที่หาไม่เจอใน catalog", B2_ROWS),
        ("B3", "Dispute Order — WeeeT แจ้งข้อพิพาท (ของไม่ตรง/ชำรุด)", B3_ROWS),
        ("B4", "Close & Rate — ยืนยันรับของ + ให้คะแนน seller", B4_ROWS),
        ("B5", "Parts in Job Context (T-36) — ซื้ออะไหล่ระหว่างงานซ่อม", B5_ROWS),
    ]

    for case_key, case_desc, rows in CASES:
        add_section_heading(doc, f"§{case_key[1:]}  {case_key} — {case_desc}", 2)
        make_table(doc, rows, case_key)

    # Screen Registry
    add_section_heading(doc, "§6  Screen Registry (T-24..T-32 + T-36 · WeeeT Parts)", 2)
    make_table(doc, REGISTRY_ROWS, "reg")

    out = os.path.join(os.path.dirname(__file__), "Parts_WeeeT.docx")
    doc.save(out)
    print(f"Saved: {out}")

    d2 = Document(out)
    total_rows = sum(len(t.rows) for t in d2.tables)
    cols = [len(t.columns) for t in d2.tables]
    print(f"rows {total_rows} cols {cols}")
    bad = [c for c in cols if c != 9]
    if bad:
        print(f"ERROR: tables with wrong cols: {bad}")
        sys.exit(1)
    else:
        print("✅ All tables = 9 columns")


if __name__ == "__main__":
    build()
