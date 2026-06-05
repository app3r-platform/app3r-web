"""สร้าง WeeeT Screen Registry Word file — T-16..T-48 (33 screens)
9 columns: รหัสจอ · ชื่อจอ/หน้าที่ · มาจาก · เงื่อนไข/เคส · ไปต่อ · แอพฯที่เห็น · เคส · หมายเหตุ · ลิงก์เปิด mockup
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WEEET_PRIMARY = RGBColor(0x16, 0x96, 0xF9)  # #1696F9

HEADERS = [
    "รหัสจอ",
    "ชื่อจอ / หน้าที่",
    "มาจาก",
    "เงื่อนไข / เคส",
    "ไปต่อ",
    "แอพฯที่เห็น",
    "เคส",
    "หมายเหตุ",
    "ลิงก์เปิด mockup\n(localhost:3003)",
]

# Base URL
BASE = "http://localhost:3003"

SCREENS = [
    # T-16 ─────────────────────────────────────────────────────────────────
    ("T-16", "รายละเอียดประกาศบริการ",
     "T-17 รายการประกาศ",
     "state D83: published / has_offer / matched / completed / cancelled\nescrow phase: none / held / released / refunded",
     "matched→completed (ยืนยันส่งมอบ ปล่อย escrow)\nmatched→cancelled (ยกเลิก คืนเงิน)\nT-17 กลับรายการ",
     "WeeeT (ช่าง)",
     "S-TA1 ส่งมอบ release escrow\nS-TA2 ยกเลิก refund",
     "Escrow Gold Point; D83 state machine; transitionListing()",
     f"{BASE}/listings/L-001"),
    # T-17 ─────────────────────────────────────────────────────────────────
    ("T-17", "ประกาศบริการของฉัน",
     "Bottom Nav / T-18 Dashboard",
     "GR-10 near-me filter (radius 20 km)\nfilter by state",
     "T-16 รายละเอียด",
     "WeeeT (ช่าง)",
     "ดูรายการ\nกรองใกล้ฉัน GR-10",
     "NearMeFilter (@app3r/ui) · roleTheme weeet-primary",
     f"{BASE}/listings"),
    # T-18 ─────────────────────────────────────────────────────────────────
    ("T-18", "Dashboard ช่าง (หน้าหลัก)",
     "Login success / Middleware redirect",
     "ช่างล็อกอินแล้ว",
     "T-01 /jobs · T-17 /listings · T-19 /today · T-12 /profile",
     "WeeeT (ช่าง)",
     "แสดงสถิติงานวันนี้",
     "สวัสดี + ชื่อช่าง + stats: งานวันนี้ / เสร็จ / งานรอ · mock data",
     f"{BASE}/dashboard"),
    # T-19 ─────────────────────────────────────────────────────────────────
    ("T-19", "งานวันนี้",
     "Bottom Nav / T-18 Dashboard",
     "jobs ของ tech-001 วันนี้ (scheduledAt today)\ncurrentStage = in_progress",
     "T-37 /jobs/[id]/progress",
     "WeeeT (ช่าง)",
     "active jobs\ndone jobs",
     "localStorage seed data; DailyQueueCard component; Phase D WebSocket (future)",
     f"{BASE}/today"),
    # T-20 ─────────────────────────────────────────────────────────────────
    ("T-20", "รายงาน / สรุปผลงาน",
     "Bottom Nav",
     "period: วันนี้ / สัปดาห์นี้ / เดือนนี้",
     "—",
     "WeeeT (ช่าง)",
     "สรุปงาน: total / completed / in_progress / assigned",
     "mock data; completion rate bar; recent jobs list; period selector",
     f"{BASE}/reports"),
    # T-21 ─────────────────────────────────────────────────────────────────
    ("T-21", "ตั้งค่า",
     "Bottom Nav / T-12 Profile",
     "ช่างล็อกอินแล้ว",
     "T-47 Logout → /login\nT-48 change-password",
     "WeeeT (ช่าง)",
     "GPS toggle\nnotification toggle\nlanguage th/en\nlogout\nเปลี่ยนรหัสผ่าน",
     "LocationPickerSection; PushPermissionBanner; changePassword()",
     f"{BASE}/settings"),
    # T-22 ─────────────────────────────────────────────────────────────────
    ("T-22", "รายการงานรับซาก",
     "Bottom Nav",
     "filter: ทั้งหมด / รับงานใหม่ / ถึงแล้ว / no-show / เสร็จสิ้น",
     "T-23 รายละเอียดงาน",
     "WeeeT (ช่างรับซาก)",
     "no-show alert badge (S9)\nactive count badge",
     "S9 no-show banner; S12 sourceRepairJobId badge · mock 3 jobs",
     f"{BASE}/scrap"),
    # T-23 ─────────────────────────────────────────────────────────────────
    ("T-23", "รับซาก — ตรวจสอบ + ดำเนินการ",
     "T-22 รายการงานรับซาก",
     "status: assigned / traveling / arrived / verifying / mismatch_reported / pickup_confirmed / no_show / completed / cancelled",
     "✅ ตรงตาม → pickup_confirmed\n⚠️ ไม่ตรง (S8) → mismatch_reported\n🚫 ไม่พบ (S9) → no_show\nT-22 กลับรายการ",
     "WeeeT (ช่างรับซาก)",
     "S8 ซากไม่ตรงสเปก\nS9 ไม่พบลูกค้า/ของ\nS12 repair badge",
     "GPS check-in; หลักฐานรูป S8; ราคาเสนอใหม่ Gold Point; no-show modal",
     f"{BASE}/scrap/SPJ-001"),
    # T-24 ─────────────────────────────────────────────────────────────────
    ("T-24", "อะไหล่ / วัสดุ (คลังร้าน)",
     "Bottom Nav",
     "search ชื่อ/SKU\nfilter category",
     "T-25 รายละเอียดอะไหล่\nT-30 ออเดอร์ของฉัน",
     "WeeeT (ช่าง)",
     "ดูคลัง\nค้นหาอะไหล่",
     "partsApi.list(); stock qty; condition: new/used/refurbished; reservedQty",
     f"{BASE}/parts"),
    # T-25 ─────────────────────────────────────────────────────────────────
    ("T-25", "รายละเอียดอะไหล่",
     "T-24 /parts",
     "part ID param",
     "T-28 cart\nT-36 /jobs/[id]/parts",
     "WeeeT (ช่าง)",
     "ดูรายละเอียด\nเพิ่มตะกร้า",
     "B5 อะไหล่; partsApi.get(id)",
     f"{BASE}/parts/p001"),
    # T-26 ─────────────────────────────────────────────────────────────────
    ("T-26", "แคตตาล็อกอะไหล่",
     "T-24 /parts",
     "—",
     "T-27 รายละเอียดในแคตตาล็อก",
     "WeeeT (ช่าง)",
     "เรียกดูแคตตาล็อก",
     "B5 รายการสั่งซื้อจากผู้ขาย",
     f"{BASE}/parts/catalog"),
    # T-27 ─────────────────────────────────────────────────────────────────
    ("T-27", "รายละเอียดในแคตตาล็อก",
     "T-26 /parts/catalog",
     "catalog item ID",
     "T-28 cart (เพิ่มตะกร้า)",
     "WeeeT (ช่าง)",
     "ดูรายละเอียดสินค้า\nเพิ่มตะกร้า",
     "B5 อะไหล่; catalog item detail",
     f"{BASE}/parts/catalog/CAT-001"),
    # T-28 ─────────────────────────────────────────────────────────────────
    ("T-28", "ตะกร้าอะไหล่",
     "T-27 catalog/[id]\nT-25 parts/[id]",
     "items in cart",
     "T-29 checkout",
     "WeeeT (ช่าง)",
     "ดูตะกร้า\nแก้จำนวน\nลบ",
     "localStorage cart; B5",
     f"{BASE}/parts/cart"),
    # T-29 ─────────────────────────────────────────────────────────────────
    ("T-29", "ชำระเงินสั่งซื้ออะไหล่",
     "T-28 cart",
     "cart items → order",
     "T-30 orders (order confirmed)",
     "WeeeT (ช่าง)",
     "ยืนยันสั่งซื้อ",
     "B5; สร้าง order; localStorage order IDs",
     f"{BASE}/parts/checkout"),
    # T-30 ─────────────────────────────────────────────────────────────────
    ("T-30", "ออเดอร์อะไหล่ของฉัน",
     "T-24 /parts\nT-29 checkout",
     "localStorage order IDs",
     "T-31 รายละเอียดออเดอร์",
     "WeeeT (ช่าง)",
     "ดูรายการออเดอร์",
     "B5; weeet_part_order_ids localStorage",
     f"{BASE}/parts/orders"),
    # T-31 ─────────────────────────────────────────────────────────────────
    ("T-31", "รายละเอียดออเดอร์อะไหล่",
     "T-30 /parts/orders",
     "order ID param",
     "T-30 กลับรายการ",
     "WeeeT (ช่าง)",
     "ดูสถานะ\nดูรายละเอียด",
     "B5 order detail; partsApi.getOrder()",
     f"{BASE}/parts/orders/ORD-001"),
    # T-32 ─────────────────────────────────────────────────────────────────
    ("T-32", "คำขออะไหล่ (Purchase Requests)",
     "T-24 /parts\nT-36 /jobs/[id]/parts",
     "purchase_request status",
     "—",
     "WeeeT (ช่าง)",
     "ดูรายการ purchase request",
     "B5; TODO BE: GET /api/v1/parts/purchase-requests/",
     f"{BASE}/parts/requests"),
    # T-33 ─────────────────────────────────────────────────────────────────
    ("T-33", "บันทึกถึงที่ (ซ่อม)",
     "T-11 job detail\nT-34 ออกเดินทาง",
     "ถ่ายรูป ≥2 (จุดสังเกต + กับลูกค้า)\nสูงสุด 3 ใบ · ≤3MB/ใบ",
     "T-11 /jobs/[id] (success)",
     "WeeeT (ช่าง)",
     "บังคับรูป 2 ใบ",
     "repairApi.arrive(); arrival_photos FormData; capture=environment",
     f"{BASE}/jobs/REP-001/arrive"),
    # T-34 ─────────────────────────────────────────────────────────────────
    ("T-34", "ออกเดินทาง (ซ่อม)",
     "T-11 job detail",
     "ต้องระบุ GPS (Geolocation API)",
     "T-11 /jobs/[id]",
     "WeeeT (ช่าง)",
     "GPS lat/lng required",
     "repairApi.depart(); departure_location {lat, lng}",
     f"{BASE}/jobs/REP-001/depart"),
    # T-35 ─────────────────────────────────────────────────────────────────
    ("T-35", "ถ่ายรูป / วิดีโอ",
     "T-37 /progress\nT-11 job detail",
     "ก่อนซ่อม / หลังซ่อม toggle",
     "กลับ (router.back)",
     "WeeeT (ช่าง)",
     "before toggle\nafter toggle\ncaption",
     "FileReader; ไม่ call API (mockup storage); capture=environment",
     f"{BASE}/jobs/REP-001/photo"),
    # T-36 ─────────────────────────────────────────────────────────────────
    ("T-36", "ขออะไหล่ (B5 สำรอง/สั่งซื้อ)",
     "T-37 /progress\nT-11 job detail",
     "source_type=warehouse: stock>0→สำรอง, stock=0→สั่งซื้อ\nsource_type=market: สั่งซื้อเสมอ",
     "submitted → submitted_reserve หรือ submitted_purchase (success state)",
     "WeeeT (ช่าง)",
     "B5 reserve (warehouse)\nB5 purchase request (market/stock=0)",
     "mockup no API; TODO BE: POST /reserve/ + /purchase-request/",
     f"{BASE}/jobs/REP-001/parts"),
    # T-37 ─────────────────────────────────────────────────────────────────
    ("T-37", "ติดตามงาน / อัปเดต",
     "T-19 /today\nT-11 job detail\nT-18 Dashboard",
     "ServiceProgress localStorage; currentStage; sub-stages",
     "T-35 /photo\nT-36 /parts\nT-11 กลับ",
     "WeeeT (ช่าง)",
     "C7 หยุดงานชั่วคราว\nC9 Dispute evidence\nprogress timeline\nstage update",
     "ServiceProgressTimeline; StepUpdateWizard; ProgressUpdateForm; evidencePhotos+GPS",
     f"{BASE}/jobs/REP-001/progress"),
    # T-38 ─────────────────────────────────────────────────────────────────
    ("T-38", "บันทึกหลังซ่อม",
     "T-37 /progress\nT-11 job detail",
     "รูปหลังซ่อม ≥3 · คลิป ≥1 (บังคับ)\nparts_used list",
     "T-11 /jobs/[id] (success)",
     "WeeeT (ช่าง)",
     "รูป before/after\nparts_used list",
     "repairApi.postRepair(); MIN_CLIPS=1; complete_photos+clips FormData",
     f"{BASE}/jobs/REP-001/post-repair"),
    # T-39 ─────────────────────────────────────────────────────────────────
    ("T-39", "ตรวจสภาพก่อนล้าง (M3)",
     "T-40 arrive (maintain)",
     "mode: normal / risk_form / risk_submitted / noshow_form / noshow_submitted",
     "normal → T-41 checklist\nrisk → WeeeR+WeeeU joint decision\nnoshow → T-46 / settled",
     "WeeeT (ช่างล้างเครื่อง)",
     "M3 ตรวจ\nM7 ไม่มาตามนัด\nD-Maintain-1 risk\nD-Maintain-2 convert-to-repair",
     "TODO BE C-4.1b: POST risk/ + convert-to-repair/ + no-show/",
     f"{BASE}/maintain/MAINT-001/inspect"),
    # T-40 ─────────────────────────────────────────────────────────────────
    ("T-40", "บันทึกถึงที่ ล้างเครื่อง (M2)",
     "T-43 depart (maintain)",
     "ถ่ายรูป ≥2 ก่อนล้าง · ≤3MB/ใบ",
     "T-39 inspect",
     "WeeeT (ช่างล้างเครื่อง)",
     "arrival_photos ≥2",
     "maintainApi.arrive(); arrival_photos FormData; เห็นสภาพก่อนล้าง",
     f"{BASE}/maintain/MAINT-001/arrive"),
    # T-41 ─────────────────────────────────────────────────────────────────
    ("T-41", "ล้างเครื่อง — Checklist (M4)",
     "T-39 inspect",
     "?type=general|deep|sanitize (searchParams)\nMaintain_CHECKLIST_ITEMS per type",
     "T-42 complete\nM9 terminate → WeeeR",
     "WeeeT (ช่างล้างเครื่อง)",
     "M4 ล้างตาม type\nM9 terminate (WeeeU ยุติงาน)",
     "MAINTAIN_CHECKLIST_ITEMS; partsApi.list() modal; terminate modal; maintainApi.complete(id)",
     f"{BASE}/maintain/MAINT-001/checklist"),
    # T-42 ─────────────────────────────────────────────────────────────────
    ("T-42", "เสร็จงาน ล้างเครื่อง (M5)",
     "T-41 checklist",
     "รูปหลังล้าง ≥2 + ลายเซ็นลูกค้า (บังคับ)",
     "T-01 /jobs (success · router.replace)",
     "WeeeT (ช่างล้างเครื่อง)",
     "M5 เสร็จ + sign\nSignaturePad",
     "SignaturePad (@/components/SignaturePad); customer_signature; maintainApi.complete()",
     f"{BASE}/maintain/MAINT-001/complete"),
    # T-43 ─────────────────────────────────────────────────────────────────
    ("T-43", "ออกเดินทาง ล้างเครื่อง (M1)",
     "T-01 /jobs หรือ Maintain job list",
     "GPS required (Geolocation API)",
     "T-40 arrive (maintain)\nT-01 กลับ",
     "WeeeT (ช่างล้างเครื่อง)",
     "M1 GPS departure_location",
     "maintainApi.depart(); {lat, lng}; background: weeet-primary/10",
     f"{BASE}/maintain/MAINT-001/depart"),
    # T-44 ─────────────────────────────────────────────────────────────────
    ("T-44", "ส่งคืนเครื่อง (Delivery flow)",
     "T-11 job detail (repair → delivery)",
     "ป้าย: en-route / completed / receipt\n(T-44 = ฐาน; ไม่มี base route ตรงๆ)",
     "completed → success\nreceipt → ลายเซ็นลูกค้า\nT-11 กลับ",
     "WeeeT (ช่าง)",
     "delivery en-route\ndelivery completed\ndelivery receipt",
     "3 sub-routes: /delivery/en-route · /delivery/completed · /delivery/receipt",
     f"{BASE}/jobs/REP-001/delivery/en-route"),
    # T-45 ─────────────────────────────────────────────────────────────────
    ("T-45", "ส่งซ่อมแบบพัสดุ (Parcel flow)",
     "T-11 job detail (scrap/parcel path)",
     "ป้าย: in-progress / tested\n(T-45 = ฐาน; ไม่มี base route ตรงๆ)",
     "tested → T-11 กลับ\nในแต่ละ sub-step ส่ง update",
     "WeeeT (ช่าง)",
     "parcel in-progress\nparcel tested",
     "2 sub-routes: /parcel/in-progress · /parcel/tested",
     f"{BASE}/jobs/REP-001/parcel/in-progress"),
    # T-46 ─────────────────────────────────────────────────────────────────
    ("T-46", "ถึงหน้างาน — ลูกค้าไม่อยู่/ไม่มาตามนัด (M7)",
     "T-39 inspect (mode = noshow_form)",
     "dev mockup artifact\npath มี /mockup segment → Phase 4 normalize ค่อยเปลี่ยน",
     "settle → T-01 /jobs",
     "WeeeT (ช่างล้างเครื่อง)",
     "M7 No-show\nลูกค้าไม่อยู่",
     "ScreenBadge T-46 label=ไม่มาตามนัด; path: /maintain/[id]/mockup/m7-noshow",
     f"{BASE}/maintain/MAINT-001/mockup/m7-noshow"),
    # T-47 ─────────────────────────────────────────────────────────────────
    ("T-47", "เข้าสู่ระบบ (Login)",
     "Middleware redirect (unauthenticated)\nLogout จาก T-21",
     "ยังไม่ล็อกอิน หรือ session หมดอายุ",
     "T-18 Dashboard (success)\nT-48 change-password (first-login)",
     "WeeeT (ช่าง)",
     "login form\nJWT auth\nfirst-login redirect",
     "นอก group (app); Auth flow; useAuth; middleware.ts",
     f"{BASE}/login"),
    # T-48 ─────────────────────────────────────────────────────────────────
    ("T-48", "เปลี่ยนรหัสผ่านครั้งแรก",
     "T-47 Login (must_change_password = true)",
     "first-login flag จาก backend",
     "T-18 Dashboard (success)",
     "WeeeT (ช่าง)",
     "first login\nforced password change",
     "Auth flow; Middleware check; changePassword(); นอก group (app)",
     f"{BASE}/change-password-first"),
]

assert len(SCREENS) == 33, f"Expected 33 screens, got {len(SCREENS)}"

# ── Build document ──────────────────────────────────────────────────────────
doc = Document()

# Page setup: landscape A4
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.0)
section.right_margin = Cm(1.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("WeeeT — รายการหน้าจอ T-16 ถึง T-48 (33 จอ)")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = WEEET_PRIMARY

doc.add_paragraph(
    "แอพ WeeeT (ช่าง) · Port 3003 · สีหลัก #1696F9 · "
    "P2+P3 Round-2 reconcile · base origin/main da0e3a9"
).runs[0].font.size = Pt(9)

doc.add_paragraph()  # spacer

# Table
tbl = doc.add_table(rows=1, cols=9)
tbl.style = "Table Grid"

# Column widths (cm, total ~27.7 cm in landscape with margins)
COL_WIDTHS = [1.4, 4.0, 3.5, 4.5, 4.0, 2.5, 3.2, 3.5, 3.0]

# Header row
hdr_row = tbl.rows[0]
for i, (cell, hdr, w) in enumerate(zip(hdr_row.cells, HEADERS, COL_WIDTHS)):
    cell.width = Cm(w)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hdr)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Blue background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1696F9")
    tcPr.append(shd)

# Data rows
for row_idx, row_data in enumerate(SCREENS):
    row = tbl.add_row()
    fill = "E8F4FE" if row_idx % 2 == 0 else "FFFFFF"
    for i, (cell, val, w) in enumerate(zip(row.cells, row_data, COL_WIDTHS)):
        cell.width = Cm(w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(8)
        # Bold screen ID column
        if i == 0:
            run.bold = True
            run.font.color.rgb = WEEET_PRIMARY
        # Alternate row shading
        if fill != "FFFFFF":
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

out_path = r"D:\ClaudeCode\App3R\wt-p2p3-weeet\WeeeT_Screens_T16-T48.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Rows: {len(tbl.rows)} (1 header + {len(tbl.rows)-1} data)")
