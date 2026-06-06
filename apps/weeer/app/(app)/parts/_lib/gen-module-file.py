r"""
P3 Module-File Generator -- App3R Parts (WeeeR B2B)
สร้างเอกสาร 9-column ครอบคลุม P1-P12 (HUB Gen 52 CMD 6)
รัน: python gen-module-file.py
ผลลัพธ์: parts-module-file-p1-p12.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Config ──────────────────────────────────────────────────────────────────
BASE_URL = "http://localhost:3001"
OUTPUT = os.path.join(os.path.dirname(__file__), "..", "..", "..", "..", "..", "..",
                      "parts-module-file-p1-p12.docx")
OUTPUT = os.path.normpath(OUTPUT)

# ── Columns ─────────────────────────────────────────────────────────────────
HEADERS = [
    "รหัสจอ",          # 1
    "ชื่อจอ / หน้าที่",  # 2
    "มาจาก (§5)",      # 3
    "เงื่อนไข / เคส",   # 4  §7
    "ไปต่อ (§6)",       # 5
    "แอพฯ / บทบาทที่เห็น (§8)",  # 6
    "เคส",             # 7
    "หมายเหตุ (Tier-1 lens)",      # 8
    "ลิงก์ mockup",    # 9
]

# ── Row data structure ───────────────────────────────────────────────────────
# (screen_id, name_role, from_screen, condition, goto, cross_app, case_code, notes, mockup_route)
ROWS = [
    # ── P1: ลงขายอะไหล่ใหม่ ──────────────────────────────────────────────
    (
        "R-29",
        "My Listings\n(มุมผู้ขาย: รายการ + ลงขายใหม่)",
        "R-51 Parts Hub (nav ด้านล่าง)",
        "เปิดหน้าขายของฉัน → กดปุ่ม + ลงขายใหม่",
        "→ R-40 (ลงขายสำเร็จ)\nผ่าน PartListingForm modal",
        "WeeeR (ร้านผู้ขาย)\n[§8] listing ใหม่ปรากฏใน R-30 ทันที",
        "P1",
        "A8: flow→R-40 success page\nA3: back nav พร้อม\nA7: listing ตัวอย่าง: คอมเพรสเซอร์ Mitsubishi 9000BTU · 800 pts",
        "/parts/my-listings",
    ),
    (
        "R-40",
        "ลงขายอะไหล่สำเร็จ\n(PARTS-NEW-SUCCESS)",
        "R-29 (+ ลงขายใหม่ → PartListingForm)",
        "listing สร้างสำเร็จ → redirect มาที่ success page",
        "→ R-29 (ดูรายการขายของฉัน)\nCTA button",
        "WeeeR (ร้านผู้ขาย)\n[§8] ผู้ซื้อทุกร้านเห็น listing ใหม่ใน R-30",
        "P1",
        "A1: ทุก action มีหน้าสำเร็จ\nA8: flow→success page\nREF หมายเลขอ้างอิงแสดงให้ติดตาม",
        "/parts/new/success",
    ),

    # ── P2: แก้ไข/ปิดรายการ ──────────────────────────────────────────────
    (
        "R-29",
        "My Listings\n(รายการของฉัน tab)",
        "R-51 Parts Hub / R-40 CTA",
        "เปิดจัดการ listing ของตน (tab: รายการของฉัน)",
        "→ R-29c (Listing Detail)\nคลิก listing card",
        "WeeeR (ร้านผู้ขาย)",
        "P2",
        "B2: label ภาษาไทยชัดเจน\nA3: tab navigation พร้อม\nA7: listing ตัวอย่าง 3 รายการ",
        "/parts/my-listings",
    ),
    (
        "R-29c",
        "Listing Detail\n(PARTS-LISTING-DETAIL)",
        "R-29 (คลิก listing card)",
        "ดูรายละเอียด listing ที่ตนเองลงขาย",
        "← back → R-29\n(edit/close actions inline)",
        "WeeeR (ร้านผู้ขาย)",
        "P2",
        "D1: media fallback ถ้ารูปโหลดไม่ได้\nB1: ชื่อสภาพ Thai: ใหม่/มือสอง/Refurb\nB2: ราคา pts แสดงชัด",
        "/parts/my-listings/mock-001",
    ),

    # ── P3: ค้นหา + กรอง ──────────────────────────────────────────────────
    (
        "R-30",
        "ตลาดอะไหล่ B2B\n(PARTS-MARKETPLACE)",
        "R-51 Parts Hub (shortcut ตลาด B2B)",
        "browse ตลาด B2B ระหว่างร้าน → ค้นหา/กรองรายการ",
        "→ R-30c (Item Detail)\nคลิก PartCard",
        "WeeeR (ร้านผู้ซื้อ)\n[§8] WeeeR (ผู้ขาย): listing ปรากฏใน marketplace",
        "P3",
        "B2: ป้าย Thai: 'ตลาดอะไหล่ B2B'\nA3: filter panel toggle\nA7: 3 listings ตัวอย่าง · 2 ร้าน\nMarketplaceStatsCard: stats realtime",
        "/parts/marketplace",
    ),
    (
        "R-30c",
        "รายละเอียดอะไหล่\n(PARTS-ITEM-DETAIL — ดูก่อนซื้อ)",
        "R-30 (Marketplace)",
        "ดูข้อมูลก่อนตัดสินใจซื้อ (stock > 0 หรือ = 0)",
        "← กลับตลาด → R-30\n(ไม่ซื้อ กลับ browse)",
        "WeeeR (ร้านผู้ซื้อ)",
        "P3",
        "D1: PartImageGallery fallback\nB1: ชื่ออะไหล่/ยี่ห้อ Thai\nB2: ราคา pts · ฿ (reference เท่านั้น)\nListingEngagement: รีวิว/ถาม-ตอบ",
        "/parts/marketplace/mock-001",
    ),

    # ── P4: สั่งซื้อปกติ ───────────────────────────────────────────────────
    (
        "R-30c",
        "รายละเอียดอะไหล่\n(stock > 0 · Gold เพียงพอ)",
        "R-30 (Marketplace)",
        "stock > 0 · balance ≥ pricePoints · ไม่ใช่ร้านตนเอง",
        "🛒 สั่งซื้อ → PlaceOrderModal\n→ escrow lock → R-33 (My Orders)",
        "WeeeR (ร้านผู้ซื้อ)\n[§8] WeeeR (ผู้ขาย) R-29 incoming tab: order ใหม่เข้า",
        "P4",
        "B3: escrow แสดงเป็น pts จำนวน (ไม่ใช่ %)\nA8: สั่งซื้อสำเร็จ → R-33\nA7: คอมเพรสเซอร์ 2 ชิ้น · 1,600 pts · escrow lock 1,600 pts",
        "/parts/marketplace/mock-001",
    ),
    (
        "R-33",
        "My Orders (buyer tab)\n(PARTS-MY-ORDERS)",
        "R-30c (order placed สำเร็จ)",
        "order ใหม่ stage=ordered ปรากฏใน buyer tab",
        "→ R-34 (Buyer Order Detail)\nคลิก OrderCard",
        "WeeeR (ร้านผู้ซื้อ)\n[§8] WeeeR (ผู้ขาย) R-29 incoming tab update",
        "P4",
        "A1: order success → กลับ R-33 รายการออเดอร์\nPointsBalanceCard: escrow held แสดง\nA7: 2 orders ตัวอย่าง (ordered/confirmed)",
        "/parts/my-orders",
    ),

    # ── P5: ผู้ขายรับออเดอร์ ─────────────────────────────────────────────
    (
        "R-29",
        "My Listings (tab: คำสั่งซื้อ)\n(มุมผู้ขาย: incoming orders)",
        "R-29 (switch tab จาก 'รายการของฉัน')",
        "order ใหม่เข้า (stage=ordered) · ผู้ขายต้องรับหรือปฏิเสธ",
        "✅ รับออเดอร์ → Confirm modal\n→ stage: ordered→confirmed",
        "WeeeR (ร้านผู้ขาย)\n[§8] WeeeR (ผู้ซื้อ) R-34: status เปลี่ยนเป็น 'ผู้ขายรับแล้ว'",
        "P5",
        "B1: ชื่อ stage Thai: 'รอผู้ขายรับ' → 'ผู้ขายรับแล้ว'\nA7: order คอมเพรสเซอร์ 2 ชิ้น · S001→S002",
        "/parts/my-listings",
    ),

    # ── P6: ผู้ขายส่งสินค้า ───────────────────────────────────────────────
    (
        "R-29",
        "My Listings (tab: คำสั่งซื้อ)\n(มุมผู้ขาย: confirmed orders)",
        "R-29 (stage=confirmed หลัง P5)",
        "order confirmed แล้ว → กดส่งสินค้า",
        "📦 ส่งสินค้า → ShipOrderModal\n(ใส่ tracking number) → stage: shipped",
        "WeeeR (ร้านผู้ขาย)\n[§8] WeeeR (ผู้ซื้อ) R-34: tracking แสดง + 'รอรับสินค้า'",
        "P6",
        "A7: tracking KE1234567890 · Kerry Express\nA8: ship success → reload R-29",
        "/parts/my-listings",
    ),

    # ── P7: ผู้ซื้อรับสินค้า ─────────────────────────────────────────────
    (
        "R-34",
        "Buyer Order Detail\n(PARTS-BUYER-ORDER — stage=shipped)",
        "R-33 (My Orders, buyer tab)",
        "stage=shipped · tracking มี · รอผู้ซื้อยืนยันรับ",
        "✅ ยืนยันรับสินค้า (checklist)\n→ ConfirmReceiveModal → stage: received\n→ InventoryImportPrompt (D-6)",
        "WeeeR (ร้านผู้ซื้อ)\n[§8] WeeeR (ผู้ขาย) R-33: escrow release → pts เข้าร้าน",
        "P7",
        "A1: receive → success (InventoryImportPrompt)\nOrderStageStepper แสดง progress\nA7: checklist 3 ข้อ · tracking KE1234567890\nD-6: ถามนำเข้าสต็อก",
        "/parts/my-orders/mock-001",
    ),

    # ── P8: ผู้ซื้อยกเลิก ─────────────────────────────────────────────────
    (
        "R-34",
        "Buyer Order Detail\n(PARTS-BUYER-ORDER — stage=ordered)",
        "R-33 (My Orders, buyer tab)",
        "stage=ordered เท่านั้น (ก่อนผู้ขาย confirm)\nมีปุ่มยกเลิกให้ผู้ซื้อ",
        "❌ ยกเลิก → CancelOrderConfirm\n→ escrow refund → listing stock คืน",
        "WeeeR (ร้านผู้ซื้อ)\n[§8] WeeeR (ผู้ขาย) R-29: stock คืน + order หายออกจาก incoming",
        "P8",
        "A7: เหตุผลยกเลิก 'สั่งผิดรายการ'\nescrow refund ทันที (ไม่ต้องรออนุมัติ)\nA1: cancel → success flash",
        "/parts/my-orders/mock-002",
    ),

    # ── P9: ผู้ขายปฏิเสธ/ยกเลิก ─────────────────────────────────────────
    (
        "R-29",
        "My Listings (tab: คำสั่งซื้อ)\n(มุมผู้ขาย: ยกเลิก order)",
        "R-29 (incoming tab · order stage=ordered/confirmed)",
        "ผู้ขายต้องการปฏิเสธ/ยกเลิก order",
        "❌ ยกเลิก → CancelOrderConfirm\n→ escrow refund → listing stock คืน",
        "WeeeR (ร้านผู้ขาย)\n[§8] WeeeR (ผู้ซื้อ) R-34/R-33: order cancelled",
        "P9",
        "A7: เหตุผล 'สต็อกหมด/ข้อมูลผิด'\ncancelledAt timestamp บันทึก",
        "/parts/my-listings",
    ),

    # ── P10: สต็อกหมด ─────────────────────────────────────────────────────
    (
        "R-30c",
        "รายละเอียดอะไหล่\n(stock = 0 — Out of Stock)",
        "R-30 (Marketplace)",
        "listing.stock === 0 → ซื้อไม่ได้",
        "← กลับตลาด → R-30\n🔔 แจ้งเตือนเมื่อมีสต็อก (mock toggle)",
        "WeeeR (ร้านผู้ซื้อ)",
        "P10",
        "A7: อะไหล่หมดสต็อก: คอมเพรสเซอร์ LG 12000BTU\nปุ่ม 'แจ้งเตือนเมื่อมีสินค้า' → notifySet state\nA3: กลับตลาด link ชัดเจน",
        "/parts/marketplace/mock-outofstock",
    ),

    # ── P11: Gold ไม่พอ ───────────────────────────────────────────────────
    (
        "R-30c",
        "รายละเอียดอะไหล่\n(Gold ไม่เพียงพอ — P11)",
        "R-30 (Marketplace)",
        "stock > 0 · balance < pricePoints → แสดง balance warning",
        "ลองสั่งซื้อ → orderError → banner error\n'ไม่สามารถดำเนินการได้'",
        "WeeeR (ร้านผู้ซื้อ)",
        "P11",
        "B3: escrow จำนวน pts ชัดเจน (ไม่ใช่ %)\nbalance check: 200 pts · ต้องการ 800 pts → insufficient\nA7: ตัวอย่าง: ร้าน S003 balance 200 pts · listing ราคา 800 pts",
        "/parts/marketplace/mock-pricey",
    ),

    # ── P12: สลับร้าน ─────────────────────────────────────────────────────
    (
        "ShopIdSwitcher\n(layout component)",
        "Shop Switcher\n(P12 — สลับร้านค้า WeeeR)",
        "parts/layout.tsx\n(persistent: แสดงทุกหน้าใน /parts/*)",
        "สลับ shopId → partsSync.emit('shop_switched')\n→ ทุกหน้าใน Parts module reload",
        "R-29 / R-30 / R-33 / R-34\nreload ด้วย shopId ใหม่",
        "WeeeR (ทุกร้าน: S001/S002/S003)\n[§8] ข้อมูลเปลี่ยนตาม role ของร้านใหม่",
        "P12",
        "A7: S001=ร้านเอ (1,200 pts) · S002=ร้านบี (3,500 pts) · S003=ร้านซี (200 pts)\ndisabled=true ระหว่าง modal เปิด (race condition guard)",
        "/parts (any Parts page)",
    ),
]

# ── Color helpers — use hex strings directly ─────────────────────────────────
# Format: "RRGGBB" (no #)
CASE_COLORS = {
    "P1":  "FFE0D6",   # orange light
    "P2":  "FFECD8",   # orange lighter
    "P3":  "DBEAFF",   # blue light
    "P4":  "D1FAE5",   # green light
    "P5":  "FEF3C7",   # yellow
    "P6":  "FEF3C7",
    "P7":  "D1FAE5",
    "P8":  "FEE2E2",   # red light
    "P9":  "FEE2E2",
    "P10": "F3F4F6",   # gray
    "P11": "FEE2E2",
    "P12": "EDE9FE",   # purple light
}

HEADER_HEX = "FF663A"   # WeeeR orange

def set_cell_bg(cell, hex_color: str):
    """Set cell background color. hex_color = 'RRGGBB' string (no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_text(cell, text: str, bold=False, size=8, hex_color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if hex_color:
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# ── Build document ───────────────────────────────────────────────────────────
doc = Document()

# Page layout (A4 landscape)
section = doc.sections[0]
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Title
title_para = doc.add_heading("App3R Parts — Module File (P1-P12)", level=1)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.runs[0]
title_run.font.color.rgb = RGBColor(0xFF, 0x66, 0x3A)

# Sub-title
sub = doc.add_paragraph("WeeeR B2B Marketplace · Phase C-6 · HUB Gen 52 CMD ⑥ · Base da0e3a9 · Context: localhost:3001")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(9)
sub.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.add_paragraph()  # spacer

# Table
table = doc.add_table(rows=1, cols=len(HEADERS))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Column widths (cm) — total ~26.7 cm
COL_WIDTHS = [1.8, 3.2, 2.8, 3.5, 3.0, 3.5, 1.2, 4.5, 3.2]

# Header row
hdr_cells = table.rows[0].cells
for i, (h, w) in enumerate(zip(HEADERS, COL_WIDTHS)):
    cell = hdr_cells[i]
    cell.width = Cm(w)
    set_cell_bg(cell, HEADER_HEX)
    set_cell_text(cell, h, bold=True, size=8, hex_color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

# Data rows
for row_data in ROWS:
    cells = table.add_row().cells
    case_code = row_data[6]  # "P1", "P2", etc.
    bg = CASE_COLORS.get(case_code, RGBColor(0xFF, 0xFF, 0xFF))

    for i, (text, w) in enumerate(zip(row_data, COL_WIDTHS)):
        cell = cells[i]
        cell.width = Cm(w)
        set_cell_bg(cell, bg)

        # Column 9 (index 8) = mockup link
        if i == 8:
            url = BASE_URL + text if text.startswith("/") else text
            set_cell_text(cell, url, size=7, hex_color="2363D0")
        else:
            set_cell_text(cell, text, size=8)

doc.add_paragraph()

# Legend
doc.add_heading("ตำนาน (Legend)", level=2)
legend_items = [
    "§5 มาจาก: จอก่อนหน้าใน flow",
    "§6 ไปต่อ: ปลายทางจาก action บนจอนี้",
    "§7 เงื่อนไข/เคส: state หรือ condition ที่ทำให้ถึงจอนี้",
    "§8 แอพฯ/บทบาทที่เห็น: cross-app view ณ จังหวะนี้",
    "Tier-1 lens: B1=Thai label · B2=UI Thai · B3=Escrow pts · D1=media fallback · A4=สีธีม #FF663A · A8=flow→success · A3=back nav · A1=success page · A7=sample data",
    "Mockup: WeeeR runs on localhost:3001 · P2+P3 branch feature/parts-p2p3 · base da0e3a9",
]
for item in legend_items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(8)

# Save
os.makedirs(os.path.dirname(OUTPUT), exist_ok=True) if os.path.dirname(OUTPUT) else None
doc.save(OUTPUT)
print(f"[OK] Saved: {OUTPUT}")
print(f"   Rows: {len(ROWS)} cases (P1-P12)")
print(f"   Cols: {len(HEADERS)}")
