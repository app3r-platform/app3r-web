"""
gen_website_screens_docx.py — สร้าง Website_Screens.docx
9-col table: รหัสจอ · หน้าที่ · มาจาก · เงื่อนไข/เคส · ไปต่อ · แอพฯที่เห็น · เคส · หมายเหตุ · ลิงก์ local
W-01 ~ W-24 (ครอบทุกจอ P1 registry · R4 W-24 mint)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_URL = "http://localhost:3004"

HEADERS = [
    "รหัสจอ",
    "หน้าที่",
    "มาจาก",
    "เงื่อนไข/เคส",
    "ไปต่อ",
    "แอพฯที่เห็น",
    "เคส",
    "หมายเหตุ",
    "ลิงก์ local",
]

# fmt: off
# 9 columns per row: รหัสจอ · หน้าที่ · มาจาก · เงื่อนไข/เคส · ไปต่อ · แอพฯที่เห็น · เคส · หมายเหตุ · ลิงก์ local
SCREENS = [
    (
        "W-01",
        "HOME — หน้าหลัก App3R",
        "(origin)",
        "public · ไม่มีเงื่อนไข",
        "W-02, W-03, W-04, W-05, W-06, W-07, W-09, W-11, W-13, W-15, W-17, W-18, W-20, W-24",
        "-",
        "default",
        "จอแรกของ Website (origin) · ไม่มี MockAnnoOrigin",
        f"{BASE_URL}/",
    ),
    (
        "W-02",
        "ABOUT — เกี่ยวกับ App3R",
        "W-01",
        "public",
        "W-01",
        "-",
        "default",
        "",
        f"{BASE_URL}/about",
    ),
    (
        "W-03",
        "CONTACT — ติดต่อเรา",
        "W-01, W-02",
        "public",
        "W-01",
        "-",
        "default",
        "",
        f"{BASE_URL}/contact",
    ),
    (
        "W-04",
        "DOWNLOAD — ดาวน์โหลดแอพฯ",
        "W-01",
        "public",
        "WeeeU store, WeeeR store",
        "WeeeU: หน้า download-store · WeeeR: หน้า download-store",
        "default",
        "แสดง store badge (App Store/Play Store) ของ WeeeU + WeeeR",
        f"{BASE_URL}/download",
    ),
    (
        "W-05",
        "FAQ — คำถามที่พบบ่อย",
        "W-01",
        "public",
        "W-01",
        "-",
        "default",
        "แสดง accordion FAQ · ลิงก์กลับหน้าหลัก",
        f"{BASE_URL}/faq",
    ),
    (
        "W-06",
        "LISTINGS-HUB — รวมประกาศทุกหมวด",
        "W-01",
        "public",
        "W-07, W-09, W-11, W-13",
        "-",
        "default",
        "Hub รวมหมวด Repair / Maintain / Resell / Scrap",
        f"{BASE_URL}/listings",
    ),
    (
        "W-07",
        "LISTINGS-REPAIR — ประกาศซ่อม",
        "W-01, W-06",
        "public",
        "W-08",
        "WeeeR: รายการรับงานซ่อม",
        "default",
        "",
        f"{BASE_URL}/listings/repair",
    ),
    (
        "W-08",
        "LISTING-REPAIR-DETAIL — รายละเอียดประกาศซ่อม",
        "W-07",
        "R1: buyer logged-in · R2: not-logged-in (prompt login)",
        "W-23 (ประวัติผู้ประกาศ)",
        "WeeeR: รายละเอียดงานซ่อม (R-job-detail)",
        "R1 posted, R2 not-logged-in",
        "auth-gated action · ลิงก์ไปหน้าประวัติผู้ประกาศ (W-23)",
        f"{BASE_URL}/listings/repair/[id]",
    ),
    (
        "W-09",
        "LISTINGS-MAINTAIN — ประกาศบำรุงรักษา",
        "W-01, W-06",
        "public",
        "W-10",
        "WeeeR: รายการรับงานบำรุง",
        "default",
        "",
        f"{BASE_URL}/listings/maintain",
    ),
    (
        "W-10",
        "LISTING-MAINTAIN-DETAIL — รายละเอียดประกาศบำรุง",
        "W-09",
        "R1: buyer logged-in · R2: not-logged-in",
        "W-23 (ประวัติผู้ประกาศ)",
        "WeeeR: รายละเอียดงานบำรุง (R-job-detail)",
        "R1 posted, R2 not-logged-in",
        "auth-gated action · ลิงก์ไปหน้าประวัติผู้ประกาศ (W-23)",
        f"{BASE_URL}/listings/maintain/[id]",
    ),
    (
        "W-11",
        "LISTINGS-RESELL — ประกาศขายมือสอง",
        "W-01, W-06",
        "public",
        "W-12",
        "WeeeT: buyer view ประกาศขาย",
        "default",
        "",
        f"{BASE_URL}/listings/resell",
    ),
    (
        "W-12",
        "LISTING-RESELL-DETAIL — รายละเอียดประกาศขาย",
        "W-11",
        "R1: listing.status==active · R2: redirect → W-12b (suspended)",
        "W-23, W-12b",
        "WeeeT: ยื่นราคา (T-offer-flow) · WeeeU: เจ้าของประกาศ (U-listing-manage)",
        "R1 active, R2 suspended",
        "ลิงก์ไปประวัติผู้ประกาศ W-23 · redirect เมื่อ status==suspended → W-12b",
        f"{BASE_URL}/listings/resell/[id]",
    ),
    (
        "W-12b",
        "LISTING-RESELL-SUSPENDED — ประกาศขายถูกระงับ",
        "W-12",
        "listing.status==suspended",
        "W-11",
        "-",
        "suspended redirect",
        "redirect จาก W-12 เมื่อ status===suspended · sub-state ของ W-12",
        f"{BASE_URL}/listings/resell/[id]/suspended",
    ),
    (
        "W-13",
        "LISTINGS-SCRAP — ประกาศขายซาก",
        "W-01, W-06",
        "public",
        "W-14",
        "WeeeT: buyer view ซาก",
        "default",
        "",
        f"{BASE_URL}/listings/scrap",
    ),
    (
        "W-14",
        "LISTING-SCRAP-DETAIL — รายละเอียดประกาศซาก",
        "W-13",
        "public",
        "W-23 (ประวัติผู้ประกาศ)",
        "WeeeT: รายละเอียดซาก (T-scrap-detail)",
        "default",
        "ลิงก์ไปประวัติผู้ประกาศ W-23",
        f"{BASE_URL}/listings/scrap/[id]",
    ),
    (
        "W-15",
        "ARTICLES — บทความ/ข่าวสาร",
        "W-01",
        "public",
        "W-16",
        "-",
        "default",
        "",
        f"{BASE_URL}/articles",
    ),
    (
        "W-16",
        "ARTICLE-DETAIL — รายละเอียดบทความ",
        "W-15",
        "public",
        "W-15",
        "-",
        "default",
        "",
        f"{BASE_URL}/articles/[id]",
    ),
    (
        "W-17",
        "PRODUCTS — สินค้า/อะไหล่",
        "W-01",
        "public",
        "W-21 (redirect ต่อ)",
        "WeeeU: สินค้าที่เชื่อมกับ listing",
        "default",
        "",
        f"{BASE_URL}/products",
    ),
    (
        "W-18",
        "REGISTER-WEEER — ลงทะเบียนช่าง WeeeR",
        "W-01, W-07, W-08",
        "public (pre-registration)",
        "W-24 (terms/privacy link)",
        "WeeeU: สร้างบัญชีช่างใหม่ (U-signup) · WeeeR: ลงทะเบียน (R-register)",
        "default",
        "A4 fix: website-brand color แทน blue chrome · เชื่อมต่อ WeeeU+WeeeR",
        f"{BASE_URL}/register/weeer",
    ),
    (
        "W-19",
        "PREVIEW — ดูตัวอย่างประกาศ",
        "WeeeU owner (token)",
        "token-gated · preview token ใน URL",
        "W-07, W-09, W-11, W-13 (ลิงก์ไปหน้ารายการ)",
        "WeeeU: เจ้าของดูตัวอย่างก่อนเผยแพร่ (U-listing-preview)",
        "preview-only",
        "เข้าได้ด้วย /preview/[token] · แสดง read-only preview",
        f"{BASE_URL}/preview/[token]",
    ),
    (
        "W-20",
        "CMS-STATIC — หน้า CMS ทั่วไป",
        "W-01",
        "slug match ตาม CMS",
        "-",
        "-",
        "default",
        "catch-all [slug] · หน้าจาก CMS dynamic route",
        f"{BASE_URL}/[slug]",
    ),
    (
        "W-21",
        "PRODUCT-REDIRECT — redirect สินค้า",
        "W-17",
        "redirect ทันที (server-side)",
        "ปลายทาง redirect (WeeeU product page หรือ external)",
        "WeeeU: หน้าสินค้า",
        "redirect stub",
        "เป็น redirect stub เท่านั้น · ไม่มี JSX render · ไม่ใส่ MockAnno",
        f"{BASE_URL}/products/[id]",
    ),
    (
        "W-22",
        "LISTING-BY-ID — ค้นหาประกาศด้วย ID",
        "W-01, W-06",
        "ค้นหาด้วย listing ID ทุก type",
        "W-07/W-09/W-11/W-13 (ตาม listing.type)",
        "-",
        "meta redirect",
        "dynamic [id] route · match ทุก listing type · redirect ไปหมวดที่ถูกต้อง",
        f"{BASE_URL}/listings/[id]",
    ),
    (
        "W-23",
        "OWNER-HISTORY — ประวัติผู้ประกาศ",
        "W-08, W-10, W-12, W-14",
        "public read-only",
        "W-08, W-10, W-12, W-14 (back)",
        "WeeeU: U-owner-profile · WeeeR: R-seller-check",
        "default",
        "D1 fix: img fallback bg-gray-100 + color:transparent + aria-label",
        f"{BASE_URL}/owners/[id]",
    ),
    (
        "W-24",
        "LEGAL-CONTENT — หน้ากฎหมาย/นโยบาย",
        "W-01, W-18",
        "slug: terms / privacy / cookies / refund · ISR 300s",
        "W-01 (breadcrumb home)",
        "-",
        "R4 mint | legal slug | ISR",
        "sidebar nav 4 เอกสาร · backward-compat refund-policy→refund · /privacy + /terms redirect มาที่นี่",
        f"{BASE_URL}/legal/[slug]",
    ),
]
# fmt: on


def set_cell_bg(cell, hex_color: str):
    """Set table cell background color (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_font(cell, bold=False, size_pt=9, color_hex=None, align=None):
    """Apply font settings to every paragraph in a cell."""
    for para in cell.paragraphs:
        if align:
            para.alignment = align
        for run in para.runs:
            run.font.bold = bold
            run.font.size = Pt(size_pt)
            if color_hex:
                r, g, b = (
                    int(color_hex[0:2], 16),
                    int(color_hex[2:4], 16),
                    int(color_hex[4:6], 16),
                )
                run.font.color.rgb = RGBColor(r, g, b)


def build_doc(out_path: str):
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.0)
        section.page_width = Cm(29.7)   # A4 landscape
        section.page_height = Cm(21.0)

    # ── Title ────────────────────────────────────────────────────────────────────
    title = doc.add_heading("Website Screen Registry — W-01 ~ W-24", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0x1E, 0x9E, 0x5A)  # website-brand green

    subtitle = doc.add_paragraph("App3R Website (Port 3004) · P1 Registry · R4 W-24 mint · ครบทุกจอ 24 รายการ")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # ── Table ────────────────────────────────────────────────────────────────────
    n_rows = 1 + len(SCREENS)
    n_cols = len(HEADERS)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Try to apply a style; fall back to no style if not available
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # keep default style; borders added manually if needed

    # Column widths (A4 landscape ~27 cm usable after margins)
    col_widths = [Cm(1.5), Cm(4.0), Cm(2.5), Cm(3.5), Cm(3.5), Cm(3.5), Cm(2.5), Cm(2.5), Cm(4.0)]
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    # ── Header row ───────────────────────────────────────────────────────────────
    hdr_row = table.rows[0]
    for i, hdr_text in enumerate(HEADERS):
        cell = hdr_row.cells[i]
        cell.text = hdr_text
        set_cell_bg(cell, "1E9E5A")  # website-brand green bg
        set_cell_font(cell, bold=True, size_pt=9, color_hex="FFFFFF")

    # ── Data rows ────────────────────────────────────────────────────────────────
    for row_idx, screen_data in enumerate(SCREENS):
        row = table.rows[row_idx + 1]
        bg = "F0FDF4" if row_idx % 2 == 0 else "FFFFFF"  # alternating light green / white

        for col_idx, cell_text in enumerate(screen_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            set_cell_bg(cell, bg)

            # รหัสจอ column: bold + brand color
            if col_idx == 0:
                set_cell_font(cell, bold=True, size_pt=9, color_hex="1E6E42")
            else:
                set_cell_font(cell, bold=False, size_pt=8)

    # ── Footer note ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        "หมายเหตุ: W-12b (suspended) = sub-state ของ W-12 · "
        "W-21 = redirect stub (ไม่มี JSX) · /privacy + /terms = redirect → W-24 · "
        "MockAnno: @/components/MockAnno · render class: mock-anno"
    )
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")
    print(f"     Rows: {len(SCREENS)} screens · {len(HEADERS)} columns")


if __name__ == "__main__":
    out_dir = os.path.join(os.path.dirname(__file__), "..", "docs")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "Website_Screens.docx")
    build_doc(out_path)
