"""
WeeeR_Screens.docx generator — HUB Gen53 P3-fix
9-col: รหัส · หน้าที่ · มาจาก · เงื่อนไข/เคส · ไปต่อ · แอพฯที่เห็น · เคส · หมายเหตุ · ลิงก์ local
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_URL = "http://localhost:3001"

# ────────────────────────────────────────────────────────────────────────────
# Screen data — (id, name_th, origin, condition, nav_to, cross_app, test_cases, notes, path)
# ────────────────────────────────────────────────────────────────────────────
SCREENS = [
    # ── Auth ──────────────────────────────────────────────────────────────
    ("R-79",  "เข้าสู่ระบบ",                  "entry",       "ยังไม่ login",          "R-01, R-80",      "-",                     "login สำเร็จ · ผิดรหัส · account ไม่มี",          "auth entry",              "/login"),
    ("R-80",  "สมัครสมาชิก (ขั้นที่ 1)",        "R-79",        "ยังไม่มีบัญชี",         "R-81",            "-",                     "กรอกครบ · password อ่อนแอ",                          "signup step 1",           "/signup"),
    ("R-81",  "ประเภทธุรกิจ",                  "R-80",        "-",                    "R-82",            "-",                     "เลือก individual · company",                        "signup step 2",           "/signup/business-type"),
    ("R-82",  "ข้อมูลธุรกิจ",                   "R-81",        "-",                    "R-83",            "-",                     "กรอกชื่อ/ที่อยู่ครบ · ไม่ครบ",                    "signup step 3",           "/signup/business-info"),
    ("R-83",  "ที่ตั้งร้าน",                    "R-82",        "-",                    "R-84",            "-",                     "ค้นหา postal code · pin map",                       "signup step 4",           "/signup/shop-location"),
    ("R-84",  "บัญชีธนาคาร",                   "R-83",        "-",                    "R-85",            "-",                     "เลือกธนาคาร · กรอกเลขบัญชี",                       "signup step 5",           "/signup/bank-account"),
    ("R-85",  "อัปโหลด KYC",                  "R-84",        "-",                    "R-86",            "-",                     "อัปโหลดไฟล์ · ไม่มีไฟล์",                          "signup step 6",           "/signup/kyc-upload"),
    ("R-86",  "ยืนยันอีเมล / OTP",             "R-85",        "-",                    "R-87",            "-",                     "OTP ถูก · OTP ผิด · หมดเวลา",                      "signup step 7",           "/signup/verify"),
    ("R-87",  "รอการอนุมัติ",                   "R-86",        "-",                    "-",               "-",                     "pending · reviewing · approved · rejected",         "KYC status",              "/signup/pending-review"),

    # ── Dashboard ─────────────────────────────────────────────────────────
    ("R-01",  "หน้าหลัก WeeeR",                "R-79 (login)", "login แล้ว",           "R-02,R-47,R-64,R-66,R-70,R-35,R-36,R-37,R-50,R-51,R-73", "WeeeU, WeeeT, Admin", "ดู KPI · nav quick links · notification badge", "central hub",           "/dashboard"),

    # ── Repair ────────────────────────────────────────────────────────────
    ("R-02",  "ประกาศซ่อม — รายการ",           "R-01",        "-",                    "R-04",            "WeeeU (U-repair-list)",  "ดู list · filter · search",                          "repair module",           "/repair/announcements"),
    ("R-03",  "ยื่นข้อเสนอซ่อม",               "R-04",        "announcement open",    "R-38",            "WeeeU (U-repair-detail)","กรอกข้อเสนอ · ส่ง · cancel",                        "bid form",                "/repair/announcements/[id]/offer"),
    ("R-04",  "รายละเอียดประกาศซ่อม",           "R-02",        "-",                    "R-03",            "WeeeU (U-repair-detail)","ดูรายละเอียด · ดู map · กด bid",                    "announce detail",         "/repair/announcements/[id]"),
    ("R-38",  "ยื่นข้อเสนอซ่อมสำเร็จ ✅",       "R-03",        "bid submit success",   "R-01",            "-",                     "success state · กด back ไป dashboard",              "bid success page",        "/repair/announcements/[id]/offer/success"),
    ("R-64",  "Dashboard ซ่อม",                "R-01",        "-",                    "R-05,R-07,R-07b,R-09", "WeeeU",            "ดู KPI ซ่อม · nav to queue",                         "repair dashboard",        "/repair/dashboard"),
    ("R-05",  "คิว Walk-in",                   "R-64",        "C1 ONSITE service",    "R-06",            "-",                     "รายการงาน walk-in",                                  "C1 queue",                "/repair/walk-in/queue"),
    ("R-06",  "รายละเอียดงาน Walk-in",          "R-05",        "-",                    "sub-routes",      "-",                     "inspect · in-progress · ready · receive · abandoned","C1 job detail",           "/repair/walk-in/[id]"),
    ("R-07",  "คิวพัสดุ (Parcel)",              "R-64",        "C3 SHIPPING service",  "R-08",            "-",                     "รายการงานพัสดุ",                                    "C3 queue",                "/repair/parcel/queue"),
    ("R-07b", "คิว Pickup",                    "R-64",        "C2 PICKUP_DELIVERY",   "R-65",            "-",                     "รายการงาน pickup",                                   "C2 queue",                "/repair/pickup/queue"),
    ("R-08",  "รายละเอียดงานพัสดุ (Parcel)",    "R-07",        "-",                    "sub-routes",      "-",                     "dispatch-tech · inspect · receive · ship-back",      "C3 job detail",           "/repair/parcel/[id]"),
    ("R-09",  "รายการงานซ่อม",                  "R-01, R-64",  "-",                    "R-11",            "-",                     "filter by status · search",                          "jobs list",               "/repair/jobs"),
    ("R-10",  "มอบหมายช่าง (ซ่อม)",             "R-11",        "ต้องการมอบหมายช่าง",   "R-11",            "-",                     "เลือก WeeeT · confirm",                              "assign tech",             "/repair/jobs/[id]/assign"),
    ("R-11",  "รายละเอียดงานซ่อม",              "R-09",        "-",                    "R-10, sub-routes","WeeeT (T-job-detail)",  "approve · dispute · progress update",                "job detail",              "/repair/jobs/[id]"),
    ("R-65",  "Pickup step flow (C2)",          "R-07b",       "C2 PICKUP_DELIVERY",   "sub-steps",       "-",                     "intake · diagnose · dispatch · ready · track",       "C2 pickup multi-step",    "/repair/pickup/[id]"),

    # ── Maintain ──────────────────────────────────────────────────────────
    ("R-47",  "คิวงานบำรุงรักษา",               "R-01",        "-",                    "R-48",            "WeeeU (U-maintain-list)","ดู list ประกาศ · filter",                            "maintain announce queue", "/maintain/queue"),
    ("R-48",  "ยื่นข้อเสนองานบำรุงรักษา",        "R-47",        "announcement open",    "R-39",            "WeeeU",                 "กรอกข้อเสนอ · submit",                               "maintain offer form",     "/maintain/queue/[id]/offer"),
    ("R-39",  "ยื่นข้อเสนอบำรุงรักษาสำเร็จ ✅",  "R-48",        "bid success",          "R-01",            "-",                     "success state",                                      "maintain offer success",  "/maintain/queue/[id]/offer/success"),
    ("R-12",  "รายการงานบำรุงรักษา",             "R-01, R-47",  "-",                    "R-14",            "-",                     "filter · search jobs",                               "maintain jobs list",      "/maintain/jobs"),
    ("R-13",  "มอบหมายช่าง (บำรุง)",             "R-14",        "-",                    "R-14",            "-",                     "เลือกช่าง · confirm",                                "assign tech",             "/maintain/jobs/[id]/assign"),
    ("R-13b", "มอบหมาย WeeeT (บำรุง)",           "R-13",        "มี WeeeT หลายคน",      "R-14",            "-",                     "เลือก WeeeT เฉพาะ",                                  "assign weeet",            "/maintain/jobs/[id]/assign/weeet"),
    ("R-14",  "รายละเอียดงานบำรุงรักษา",          "R-12",        "-",                    "R-13, sub-routes","WeeeT (T-job-detail)",  "progress · withdraw · complete",                    "maintain job detail",     "/maintain/jobs/[id]"),

    # ── Resell ────────────────────────────────────────────────────────────
    ("R-66",  "Resell Hub (ขายมือสอง)",          "R-01",        "-",                    "R-15,R-15c,R-17,R-20,R-67,R-68,R-69", "-",  "ดู KPI · quick nav",                                "resell hub",              "/resell"),
    ("R-67",  "คลังสินค้ามือสอง",                "R-66",        "-",                    "R-67b, R-67c",    "-",                     "filter by status · search",                          "resell inventory",        "/resell/inventory"),
    ("R-67b", "เพิ่มสินค้ามือสอง",               "R-67",        "-",                    "R-67c",           "-",                     "กรอกข้อมูล · scan barcode · SKU lookup",             "new inventory item",      "/resell/inventory/new"),
    ("R-67c", "รายละเอียดสินค้ามือสอง",           "R-67",        "-",                    "R-15 (ประกาศขาย)","WeeeU",                 "ดูรายละเอียด · ประกาศขาย · edit",                   "inventory detail",        "/resell/inventory/[id]"),
    ("R-15",  "ประกาศขายของฉัน",                 "R-66",        "-",                    "R-16",            "WeeeU (U-resell-list)",  "ดูประกาศ · filter status",                           "my listings",             "/resell/listings"),
    ("R-16",  "รายละเอียดประกาศขาย",             "R-15",        "-",                    "-",               "WeeeU",                 "ดู offer ที่ได้รับ · suspend · re-list",             "listing detail",          "/resell/listings/[id]"),
    ("R-41",  "ประกาศขายสำเร็จ ✅",              "R-15 (new)",  "listing submit success","R-15",            "-",                     "success state",                                      "new listing success",     "/resell/listings/new/success"),
    ("R-15c", "รับซื้อมือสอง (B6)",              "R-66",        "-",                    "R-15b",           "WeeeU (sell-flow)",      "เปิดหน้าซื้อ",                                       "buy entry",               "/resell/buy"),
    ("R-15b", "B6 Wizard ตีราคา",               "R-15c",       "-",                    "R-67",            "WeeeU",                 "scan barcode · ตีราคา · บันทึก",                    "B6 price wizard",         "/resell/buy/wizard"),
    ("R-17",  "Marketplace C2C",                "R-66",        "-",                    "R-19",            "WeeeU (U-marketplace)",  "ดูรายการ · filter · search",                         "C2C list",                "/resell/marketplace"),
    ("R-19",  "รายละเอียดสินค้า C2C",            "R-17",        "-",                    "R-18",            "WeeeU",                 "ดูรายละเอียด · ยื่นข้อเสนอ",                        "C2C detail",              "/resell/marketplace/[id]"),
    ("R-18",  "ยื่นข้อเสนอซื้อ (C2C)",           "R-19",        "-",                    "-",               "-",                     "กรอกราคา · delivery method",                         "pair3 offer",             "/resell/marketplace/[id]/offer"),
    ("R-20",  "รายการซื้อของฉัน",                "R-66",        "-",                    "R-23",            "-",                     "filter by status · รายการซื้อ",                     "purchases list",          "/resell/purchases"),
    ("R-23",  "รายละเอียดการซื้อ",               "R-20",        "-",                    "R-21, R-22",      "-",                     "ดูรายละเอียด · inspection period",                  "purchase detail",         "/resell/purchases/[id]"),
    ("R-21",  "ตรวจสอบสินค้า (R1)",             "R-23",        "อยู่ใน inspection period","R-23",           "-",                     "accept · dispute",                                   "inspect action",          "/resell/purchases/[id]/inspect"),
    ("R-22",  "พิพาทการซื้อ (R8)",               "R-23",        "มีปัญหา",              "R-23",            "-",                     "กรอกเหตุผล · upload evidence",                       "dispute action",          "/resell/purchases/[id]/dispute"),
    ("R-68",  "ข้อเสนอของฉัน",                   "R-66",        "-",                    "-",               "-",                     "filter status · escrow countdown · ถอนข้อเสนอ",     "my offers",               "/resell/offers"),
    ("R-69",  "รายการซื้อขาย",                   "R-66",        "-",                    "R-69b",           "-",                     "filter status · active/disputed count",              "transactions list",       "/resell/transactions"),
    ("R-69b", "รายละเอียดซื้อขาย",               "R-69",        "-",                    "-",               "-",                     "ดูรายละเอียด transaction",                           "transaction detail",      "/resell/transactions/[id]"),
    ("R-42",  "รายละเอียด Listing (Meta)",       "R-43, R-44",  "-",                    "-",               "WeeeU",                 "ดูรายละเอียด listing กลาง",                          "meta listing",            "/listings/[id]"),

    # ── Scrap ─────────────────────────────────────────────────────────────
    ("R-70",  "Scrap Hub (ตลาดซาก)",             "R-01",        "-",                    "R-71,R-72,R-24,R-27", "-",              "ดู feed · browse · my scrap jobs",                   "scrap hub",               "/scrap"),
    ("R-71",  "รายละเอียดซากของฉัน",              "R-70",        "-",                    "R-25",            "-",                     "ดูรายละเอียด · ยื่นข้อเสนอ",                        "my scrap item",           "/scrap/[id]"),
    ("R-72",  "เลือกซื้อซาก (Browse)",            "R-70",        "-",                    "R-78, R-26",      "-",                     "filter grade/price · ดูรายการ",                     "scrap browse",            "/scrap/browse"),
    ("R-78",  "รายละเอียดซากที่ซื้อ",             "R-72",        "-",                    "R-25",            "-",                     "ดูรายละเอียด · ยื่น bid",                            "scrap browse detail",     "/scrap/browse/[id]"),
    ("R-24",  "ประกาศขายซาก (redirect)",          "R-70",        "redirect → R-72",      "R-72",            "-",                     "redirect 301 ไป /scrap/browse",                      "legacy redirect",         "/scrap/announcements"),
    ("R-26",  "รายละเอียดประกาศขายซาก",           "R-72",        "-",                    "R-25",            "WeeeU (U-scrap-detail)","ดูรายละเอียด · เปิดรับข้อเสนอ",                    "scrap announce detail",   "/scrap/announcements/[id]"),
    ("R-25",  "ยื่นข้อเสนอซื้อซาก",               "R-26",        "announcement open",    "R-27",            "WeeeU (U-scrap-detail)","เลือก mode buy/free · B3 escrow · submit",           "scrap bid form",          "/scrap/announcements/[id]/offer"),
    ("R-27",  "รายการงานซาก",                    "R-25 (post-offer), R-70", "-", "R-28",             "-",                     "ดูรายการ scrap jobs",                                "scrap jobs list",         "/scrap/jobs"),
    ("R-28",  "รายละเอียดงานซาก",                 "R-27",        "-",                    "R-28b,R-28c,R-28d,R-28e","-",           "ดู decision options",                                "scrap job detail",        "/scrap/jobs/[id]"),
    ("R-28b", "Decision: Resell as Scrap (S1)",   "R-28",        "decision S1",          "R-28",            "-",                     "ยืนยัน S1",                                          "scrap s1 decision",       "/scrap/jobs/[id]/resell-as-scrap"),
    ("R-28c", "Decision: Resell Parts (S2)",      "R-28",        "decision S2",          "R-28",            "-",                     "ยืนยัน S2 → Parts module",                           "scrap s2 decision",       "/scrap/jobs/[id]/resell-parts"),
    ("R-28d", "Decision: Repair & Sell (S3)",     "R-28",        "decision S3",          "R-28",            "-",                     "ยืนยัน S3 → Repair flow",                            "scrap s3 decision",       "/scrap/jobs/[id]/repair-and-sell"),
    ("R-28e", "Decision: Dispose (S4)",           "R-28",        "decision S4",          "R-28",            "-",                     "ยืนยัน S4 · กรอก disposal note",                    "scrap s4 decision",       "/scrap/jobs/[id]/dispose"),

    # ── Parts ─────────────────────────────────────────────────────────────
    ("R-51",  "Parts Hub (อะไหล่)",               "R-01",        "-",                    "R-52,R-53,R-54,R-57,R-58,R-29,R-30,R-31,R-33,R-60,R-61,R-62", "WeeeU (U-parts-browse)", "ดู quick nav · cross-app panel", "parts hub", "/parts"),
    ("R-52",  "Parts Dashboard",                  "R-51",        "-",                    "R-59, R-58, R-31","WeeeU",                 "ดู KPI · stock alerts",                              "parts dashboard",         "/parts/dashboard"),
    ("R-53",  "คลังอะไหล่ (Inventory)",            "R-51",        "-",                    "R-63, R-57",      "-",                     "grid/list toggle · search · filter · low-stock",    "parts inventory",         "/parts/inventory"),
    ("R-57",  "เพิ่มอะไหล่ใหม่",                  "R-53",        "-",                    "R-40",            "-",                     "กรอกข้อมูล · categories · pricing",                 "new part form",           "/parts/new"),
    ("R-40",  "เพิ่มอะไหล่สำเร็จ ✅",             "R-57",        "submit success",       "R-52",            "-",                     "success state",                                      "parts new success",       "/parts/new/success"),
    ("R-63",  "รายละเอียดอะไหล่",                 "R-53, R-58",  "-",                    "R-63b, R-63c, R-63d","-",                 "ดูรายละเอียด · stock level",                         "part detail",             "/parts/[id]"),
    ("R-63b", "แก้ไขอะไหล่",                     "R-63",        "-",                    "R-63",            "-",                     "edit fields · save",                                 "part edit",               "/parts/[id]/edit"),
    ("R-63c", "รับสต๊อกอะไหล่ (Stock-in)",        "R-63",        "-",                    "R-63",            "-",                     "กรอกจำนวน · note",                                  "stock-in",                "/parts/[id]/stock-in"),
    ("R-63d", "ปรับสต๊อกอะไหล่ (Adjust)",         "R-63",        "-",                    "R-63",            "-",                     "+/- จำนวน · note required",                          "stock adjust",            "/parts/[id]/stock-adjust"),
    ("R-54",  "ตะกร้าอะไหล่",                    "R-30 (add to cart)", "-",            "R-55",            "-",                     "ดู items · update qty · remove",                    "parts cart",              "/parts/cart"),
    ("R-55",  "Checkout อะไหล่",                  "R-54",        "-",                    "R-33b",           "-",                     "ยืนยัน order · B3 escrow",                           "parts checkout",          "/parts/checkout"),
    ("R-56",  "ถอดชิ้นส่วนจากซาก",               "R-52",        "มี scrap job S2",      "-",               "-",                     "เลือก scrap job · บันทึกชิ้นส่วน",                  "disassemble from scrap",  "/parts/disassemble"),
    ("R-58",  "จองอะไหล่",                       "R-52",        "-",                    "R-63",            "-",                     "ดูรายการ reservation",                               "parts reservations",      "/parts/reservations"),
    ("R-59",  "ความเคลื่อนไหวสต๊อก",              "R-52",        "-",                    "R-59b",           "-",                     "filter by type/date",                                "stock movements",         "/parts/movements"),
    ("R-59b", "รายละเอียด Movement",              "R-59",        "-",                    "-",               "-",                     "ดูรายละเอียด movement",                              "movement detail",         "/parts/movements/[id]"),
    ("R-60",  "กล่องข้อความขออะไหล่ (Inbox)",    "R-51",        "-",                    "-",               "WeeeU (U-parts-req)",   "ดูคำขอจาก WeeeU/WeeeR อื่น",                        "parts request inbox",     "/parts/requests/inbox"),
    ("R-61",  "คำขออะไหล่ของฉัน",               "R-51",        "-",                    "R-62",            "-",                     "ดูคำขอที่ส่งออกไป",                                  "my requests",             "/parts/requests/my"),
    ("R-62",  "ส่งคำขออะไหล่ใหม่",               "R-51, R-61",  "-",                    "-",               "-",                     "กรอกชื่ออะไหล่ · urgency · submit",                 "new request form",        "/parts/requests/new"),
    ("R-29",  "ประกาศขายอะไหล่ (My Listings)",   "R-51",        "-",                    "R-29b, R-29c",    "WeeeU (U-parts-buy)",    "ดูประกาศอะไหล่ · filter",                            "parts my listings",       "/parts/my-listings"),
    ("R-29b", "ลงประกาศขายอะไหล่ใหม่",            "R-29",        "-",                    "-",               "-",                     "กรอกรายละเอียด · กำหนดราคา",                        "new parts listing",       "/parts/my-listings/new"),
    ("R-29c", "รายละเอียดประกาศอะไหล่",           "R-29",        "-",                    "-",               "WeeeU",                 "ดูประกาศ · edit · suspend",                          "parts listing detail",    "/parts/my-listings/[id]"),
    ("R-30",  "Marketplace อะไหล่",               "R-51",        "-",                    "R-30c",           "WeeeU (U-parts-browse)","ดู feed · filter · search",                          "parts marketplace",       "/parts/marketplace"),
    ("R-30c", "รายละเอียดอะไหล่ใน Marketplace",   "R-30",        "-",                    "R-30b, R-54",     "WeeeU",                 "ดูรายละเอียด · Smart-pick · Add to cart",           "parts marketplace detail","parts/marketplace/[id]"),
    ("R-30b", "Smart-pick อะไหล่",               "R-30c",       "-",                    "R-54",            "-",                     "AI suggest · เลือก · add to cart",                  "smart picker",            "/parts/marketplace/[id]/smart-pick"),
    ("R-31",  "Orders ฝั่งขาย",                  "R-52",        "-",                    "R-32",            "-",                     "ดู orders ที่รับ · filter",                           "seller orders",           "/parts/orders"),
    ("R-32",  "รายละเอียด Order (ขาย)",           "R-31",        "-",                    "sub-routes",      "WeeeU",                 "ดูรายละเอียด · dispute · rate",                      "seller order detail",     "/parts/orders/[id]"),
    ("R-33",  "Orders ฝั่งซื้อ (My Orders)",      "R-51",        "-",                    "R-34, R-33b",     "-",                     "ดู orders ที่สั่ง · filter",                         "buyer orders",            "/parts/my-orders"),
    ("R-33b", "Order ใหม่สำเร็จ",                "R-55",        "checkout complete",    "R-33",            "-",                     "success state · กด back",                            "new order success",       "/parts/my-orders/new"),
    ("R-34",  "รายละเอียด Order (ซื้อ)",          "R-33",        "-",                    "sub-routes",      "WeeeU",                 "ดูรายละเอียด · return · warranty",                  "buyer order detail",      "/parts/my-orders/[id]"),

    # ── Jobs Board ────────────────────────────────────────────────────────
    ("R-43",  "Jobs Board — ประกาศทั้งหมด",      "R-01",        "-",                    "R-42, R-45, R-46","WeeeU (U-jobs-feed)",    "ดู feed ประกาศ · filter type",                       "jobs listings",           "/jobs/listings"),
    ("R-44",  "Jobs Board — คิวของฉัน",           "R-01",        "-",                    "R-42",            "-",                     "ดูคิวงานที่รับแล้ว",                                 "jobs queue",              "/jobs/queue"),
    ("R-45",  "Service Listings (Maintain)",      "R-43",        "-",                    "R-45b",           "WeeeU",                 "ดูรายการ maintain listings",                         "listings maintain",       "/listings/maintain"),
    ("R-45b", "รายละเอียด Maintain Listing",      "R-45",        "-",                    "R-48",            "WeeeU",                 "ดูรายละเอียด · กด bid",                              "maintain listing detail", "/listings/maintain/[id]"),
    ("R-46",  "Service Listings (Repair)",        "R-43",        "-",                    "R-46b",           "WeeeU",                 "ดูรายการ repair listings",                           "listings repair",         "/listings/repair"),
    ("R-46b", "รายละเอียด Repair Listing",        "R-46",        "-",                    "R-03",            "WeeeU",                 "ดูรายละเอียด · กด bid",                              "repair listing detail",   "/listings/repair/[id]"),

    # ── Account / Settings ────────────────────────────────────────────────
    ("R-35",  "จัดการทีมช่าง (Staff)",            "R-01",        "-",                    "R-49",            "-",                     "ดูรายชื่อช่าง · เพิ่ม WeeeT",                        "staff management",        "/staff"),
    ("R-49",  "จัดการ WeeeT",                    "R-35, R-01",  "-",                    "-",               "-",                     "ดู WeeeT ทั้งหมด · invite · remove",                 "manage technicians",      "/manage-technicians"),
    ("R-36",  "Wallet",                           "R-01",        "-",                    "R-74, R-75, R-76, R-77","-",             "ดู balance · Gold/Silver",                            "wallet hub",              "/wallet"),
    ("R-74",  "เติม Gold (Deposit)",              "R-36",        "-",                    "-",               "-",                     "เลือกจำนวน · confirm",                               "deposit",                 "/wallet/deposit"),
    ("R-75",  "ประวัติธุรกรรม",                   "R-36",        "-",                    "-",               "-",                     "filter date · type",                                 "wallet history",          "/wallet/history"),
    ("R-76",  "Settlement (รับเงิน)",             "R-36",        "-",                    "-",               "-",                     "ดู pending settlements",                             "settlements",             "/wallet/settlements"),
    ("R-77",  "ถอน Gold (Withdraw)",              "R-36",        "-",                    "-",               "-",                     "กรอกจำนวน · confirm",                               "withdraw",                "/wallet/withdraw"),
    ("R-37",  "โปรไฟล์ร้าน",                     "R-01",        "-",                    "-",               "-",                     "ดู/แก้ไขข้อมูลร้าน",                                "profile",                 "/profile"),
    ("R-73",  "บริการของร้าน (Services)",         "R-01",        "-",                    "R-73b, R-73c",    "-",                     "ดูบริการที่ลงทะเบียน · เพิ่ม",                       "services list",           "/services"),
    ("R-73b", "เพิ่มบริการใหม่",                  "R-73",        "-",                    "R-73",            "-",                     "กรอกรายละเอียด · submit",                            "new service",             "/services/new"),
    ("R-73c", "แก้ไขบริการ",                     "R-73",        "-",                    "R-73",            "-",                     "แก้ไข fields · save",                               "edit service",            "/services/[id]/edit"),
    ("R-50",  "การแจ้งเตือน",                    "R-01",        "-",                    "-",               "-",                     "mark read · ดูรายการแจ้งเตือน",                     "notifications",           "/notifications"),
]


# ────────────────────────────────────────────────────────────────────────────
# Helper: set cell background
# ────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])


def bold_para(cell, text, font_size=9):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def cell_text(cell, text, font_size=8):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text) if text else "-")
    run.font.size = Pt(font_size)


# ────────────────────────────────────────────────────────────────────────────
# Build document
# ────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page orientation: landscape
section = doc.sections[0]
section.orientation = 1  # WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# Title
title = doc.add_heading("WeeeR Screen Registry — P1 (R-01 .. R-87)", level=1)
title.runs[0].font.color.rgb = RGBColor(0xFF, 0x66, 0x3A)  # WeeeR orange

subtitle = doc.add_paragraph(f"App3R-WeeeR · port 3001 · {len(SCREENS)} จอ · HUB Gen53 P3-fix")
subtitle.runs[0].font.size = Pt(9)
subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# Table
COLS = ["รหัส", "หน้าที่", "มาจาก", "เงื่อนไข / เคส", "ไปต่อ", "แอพฯที่เห็น", "เคส", "หมายเหตุ", "ลิงก์ local"]
WIDTHS = [1.4, 4.2, 2.2, 3.0, 3.0, 2.8, 4.0, 2.5, 4.5]  # cm (landscape A4 ≈ 27cm usable)

table = doc.add_table(rows=1, cols=9)
table.style = "Table Grid"

# Header row
hdr = table.rows[0]
for i, (col, w) in enumerate(zip(COLS, WIDTHS)):
    cell = hdr.cells[i]
    set_cell_bg(cell, "FF663A")
    bold_para(cell, col, font_size=9)

# Data rows
for idx, row_data in enumerate(SCREENS):
    rid, name, origin, cond, nav, xapp, tests, notes, path = row_data
    row = table.add_row()
    bg = "FFF8F5" if idx % 2 == 0 else "FFFFFF"
    values = [rid, name, origin, cond, nav, xapp, tests, notes, BASE_URL + path]
    for i, val in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        cell_text(cell, val, font_size=8)
        # Bold screen ID
        if i == 0:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xB8, 0x30, 0x0E)

# Output path
out_dir = r"D:\ClaudeCode\App3R\App3R-System"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "WeeeR_Screens.docx")
doc.save(out_path)
print(f"✅ Saved: {out_path}")
print(f"   Screens: {len(SCREENS)}")
print(f"   Columns: {len(COLS)}")
