"""
gen_maintain_docx.py — แปลง docs/maintain-module-flow.md → docs/Maintain_Module.docx
ใช้ python-docx · 9-col table · 1 จอ = 1 แถว
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ─────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).parent.parent
MD_FILE  = ROOT / "docs" / "maintain-module-flow.md"
OUT_FILE = ROOT / "docs" / "Maintain_Module.docx"

# ─── Column definitions ─────────────────────────────────────────────────────────
HEADERS = [
    "รหัสจอ",
    "ชื่อจอ / หน้าที่",
    "มาจาก",
    "เงื่อนไข / เคส",
    "ไปต่อ",
    "แอพฯที่เห็น",
    "เคส",
    "หมายเหตุ",
    "ลิงก์ local",
]
# Column widths in cm (total ~28 cm for A4 landscape)
COL_WIDTHS = [1.8, 3.2, 2.4, 3.0, 2.4, 2.0, 1.5, 3.5, 5.2]

# ─── Colors ─────────────────────────────────────────────────────────────────────
COLOR_HEADER_BG = RGBColor(0x1E, 0x40, 0xAF)   # blue-800
COLOR_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)   # white
COLOR_SECTION_BG = RGBColor(0xDB, 0xEA, 0xFE)  # blue-100
COLOR_SECTION_FG = RGBColor(0x1E, 0x3A, 0x8A)  # blue-900
COLOR_ALT_ROW    = RGBColor(0xF8, 0xFA, 0xFF)  # very light blue
COLOR_NEW_ID     = RGBColor(0x92, 0x40, 0x0E)  # amber-800 for 🆕 rows

# ─── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))  # twips
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def add_para_text(cell, text: str, bold=False, italic=False,
                  font_size=8, color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear cell and add formatted paragraph."""
    cell.paragraphs[0].clear()
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = color


# ─── Parse MD ────────────────────────────────────────────────────────────────────
def parse_md(path: Path):
    """
    Returns list of sections:
    [{"title": "M1 — ...", "rows": [["col1","col2",...], ...]}, ...]
    """
    text = path.read_text(encoding="utf-8")
    sections = []
    current_section = None
    in_table = False

    for line in text.splitlines():
        line = line.strip()

        # Section header (## M1 ...)
        if line.startswith("## M") or line.startswith("## สรุป"):
            if current_section:
                sections.append(current_section)
            current_section = {"title": line.lstrip("# ").strip(), "rows": []}
            in_table = False
            continue

        if current_section is None:
            continue

        # Table row
        if line.startswith("|") and not line.startswith("|---"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            # Skip header row (first row with column names)
            if cells and cells[0] == "รหัสจอ":
                continue
            # Skip separator rows
            if all(set(c) <= {"-", " ", ":"} for c in cells if c):
                continue
            if len(cells) >= 9:
                current_section["rows"].append(cells[:9])
            elif len(cells) >= 5 and current_section["title"].startswith("สรุป"):
                # Summary table has fewer cols — skip for main table
                current_section["rows"].append(cells + [""] * (9 - len(cells)))

    if current_section:
        sections.append(current_section)

    return sections


# ─── Build DOCX ──────────────────────────────────────────────────────────────────
def build_docx(sections, out_path: Path):
    doc = Document()

    # Page setup — A4 Landscape
    section = doc.sections[0]
    section.page_width  = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Maintain Module — Flow Table (M1–M9)")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        "P2+P3 Maintain Gen 4 · 2026-06-06  |  "
        "WeeeU :3001  ·  WeeeR :3002  ·  WeeeT :3003  ·  Admin :3004  ·  Website :3005"
    )
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # spacer

    for sec in sections:
        if not sec["rows"]:
            continue

        # ── Section heading ─────────────────────────────────────────────────────
        h_para = doc.add_paragraph()
        h_run = h_para.add_run(sec["title"])
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = COLOR_SECTION_FG

        # ── Table ──────────────────────────────────────────────────────────────
        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        set_row_height(hdr, 0.8)
        for i, (cell, header) in enumerate(zip(hdr.cells, HEADERS)):
            cell.width = Cm(COL_WIDTHS[i])
            set_cell_bg(cell, COLOR_HEADER_BG)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_para_text(cell, header, bold=True, font_size=8,
                          color=COLOR_HEADER_FG, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Data rows
        for r_idx, row_data in enumerate(sec["rows"]):
            row = table.add_row()
            set_row_height(row, 0.6)
            is_new_id = "🆕" in (row_data[0] if row_data else "")
            alt_bg = (r_idx % 2 == 1)

            for i, (cell, text) in enumerate(zip(row.cells, row_data)):
                cell.width = Cm(COL_WIDTHS[i])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Background
                if is_new_id:
                    set_cell_bg(cell, RGBColor(0xFF, 0xF7, 0xED))  # orange-50
                elif alt_bg:
                    set_cell_bg(cell, COLOR_ALT_ROW)

                # Font color
                txt_color = COLOR_NEW_ID if is_new_id and i == 0 else None

                add_para_text(cell, text, font_size=7.5, color=txt_color)

        # Column widths (set after table creation)
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(COL_WIDTHS[i])

        doc.add_paragraph()  # spacer between sections

    # Footer note
    note = doc.add_paragraph()
    note_run = note.add_run(
        "🆕 = New IDs minted (R4 rule) — ต้อง register ใน P1 Registry  |  "
        "หมายเหตุ: IDs ที่มี '/' = sub-state ของหน้าเดิม (ไม่ใช่ route ใหม่)"
    )
    note_run.font.size = Pt(8)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"✅ Saved: {out_path}")


# ─── Main ────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"Reading: {MD_FILE}")
    sections = parse_md(MD_FILE)
    total_rows = sum(len(s["rows"]) for s in sections)
    print(f"Sections: {len(sections)}  |  Total rows: {total_rows}")
    for s in sections:
        print(f"  [{len(s['rows']):3d} rows] {s['title'][:60]}")
    build_docx(sections, OUT_FILE)
    print(f"\nDone → {OUT_FILE}")
