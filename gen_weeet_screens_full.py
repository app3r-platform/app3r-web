"""สร้าง WeeeT Screen Registry Word file — T-01..T-48 (48 screens)
EXC-1: ขยายจาก T-16..T-48 → T-01..T-48 (เพิ่ม T-07/T-13/T-14 + T-01..T-15 ทั้งหมด)
9 columns: รหัสจอ · ชื่อจอ/หน้าที่ · มาจาก · เงื่อนไข/เคส · ไปต่อ · แอพฯที่เห็น · เคส · หมายเหตุ · ลิงก์เปิด mockup
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WEEET_PRIMARY = RGBColor(0x16, 0x96, 0xF9)  # #1696F9

HEADERS = [
    "รหัสจอ",
    "ชื่อจอ / หน้าที่",
    "มาจาก",
    "เงื่อนไข / เคส",
    "ไปต่อ",
    "แอพฯที่เห็น",
    "เคส",
    "หมายเหตุ",
    "ลิงก์เปิด mockup\n(localhost:3003)",
]

BASE = "http://localhost:3003"

SCREENS = [
    # ══════════════════════════════════════════════════════════════════════════
    # T-01..T-15 — Core Jobs + Auth + Profile
    # ══════════════════════════════════════════════════════════════════════════

    # T-01 ─────────────────────────────────────────────────────────────────
    ("T-01", "รายการงาน (Job List)",
     "Bottom Nav / T-18 Dashboard",
     "tab: ทั้งหมด / on_site / pickup / parcel / done\nexclude: walk_in (ของ WeeeR)",
     "T-11 รายละเอียดงาน\nT-34 ออกเดินทาง (shortcut)",
     "WeeeT (ช่าง)",
     "ดูงานทั้งหมด\nกรองตาม service_type",
     "repairApi.list(); EXCLUDED_SERVICE_TYPES=[walk_in]; STATUS_LABEL map 12 states",
     f"{BASE}/jobs"),

    # T-02 ─────────────────────────────────────────────────────────────────
    ("T-02", "วินิจฉัย / เลือกสาขา",
     "T-11 รายละเอียดงาน /jobs/[id]",
     "branch B1.1 (ราคาเดิม) / B1.2 (ราคาใหม่+อะไหล่) / B2.1 (ยกเลิก) / B2.2 (รับซาก)\nB1.2 ต้องมีอะไหล่ ≥1",
     "T-11 /jobs/[id] (ยืนยันวินิจฉัยสำเร็จ)",
     "WeeeT (ช่าง) · WeeeR (อนุมัติ B1.2) · WeeeU (รับทราบสาขา)",
     "B1.1 ซ่อมราคาเดิม\nB1.2 ราคาใหม่+อะไหล่\nB2.1 ยกเลิก\nB2.2 รับซาก",
     "repairApi.diagnose(); DiagnosePayload; branch+notes บังคับ; B1.2→proposed_price+parts",
     f"{BASE}/jobs/REP-001/diagnose"),

    # T-03 ─────────────────────────────────────────────────────────────────
    ("T-03", "บันทึกผลซ่อม",
     "T-11 รายละเอียดงาน /jobs/[id]\nT-37 ติดตามงาน",
     "ป้าย: กำลังซ่อม (in-progress) / ทดสอบ (tested) / base\nOTP D2 pre-fill (mock, Phase 3)",
     "T-15 success (OTP confirm)\nT-11 กลับ",
     "WeeeT (ช่าง)",
     "บันทึก repairNote + parts + finalPrice\nOTP verify (D2)",
     "D2 Mock OTP pre-fill; repairNote/parts/finalPrice form; 3 sub-routes: base/in-progress/tested",
     f"{BASE}/jobs/REP-001/repair"),

    # T-04 ─────────────────────────────────────────────────────────────────
    ("T-04", "ข้อมูลงาน Pickup",
     "T-11 รายละเอียดงาน /jobs/[id]",
     "ป้าย: กำลังไปรับ (en-route) / ถึงจุดรับ (arrived) / ถึงร้าน (at-shop) / ใบรับของ (receipt)\n4 sub-routes + base",
     "T-11 กลับ\nen-route→arrived→at-shop→receipt flow",
     "WeeeT (ช่าง)",
     "Pickup en-route\nPickup arrived\nPickup at-shop\nPickup receipt",
     "Pickup sub-flow; 4 sub-routes: /pickup/en-route · /arrived · /at-shop · /receipt; GPS check",
     f"{BASE}/jobs/REP-001/pickup/en-route"),

    # T-05 ─────────────────────────────────────────────────────────────────
    ("T-05", "รออะไหล่ / นัดหมายใหม่ (C3)",
     "T-11 รายละเอียดงาน /jobs/[id]",
     "C3-SCHEDULE state: draft / awaiting / confirmed\nรออะไหล่หรือนัดลูกค้าใหม่",
     "confirmed → T-11 กลับ",
     "WeeeT (ช่าง) · WeeeR (อนุมัติ) · WeeeU (รับทราบนัด)",
     "C3 รออะไหล่\nC3 นัดหมายใหม่\nstate toggle tabs",
     "Mockup-only; state tabs draft/awaiting/confirmed; TODO BE: PATCH /jobs/[id]/schedule",
     f"{BASE}/jobs/REP-001/schedule"),

    # T-06 ─────────────────────────────────────────────────────────────────
    ("T-06", "เสนอรับซาก (Scrap Offer)",
     "T-11 รายละเอียดงาน /jobs/[id] (B2.2 path)",
     "ราคารับซาก (Gold Point)\ngold-lock 24h mock (lock icon)",
     "submitted → T-11 กลับ (scrap-offer locked)",
     "WeeeT (ช่าง) · WeeeR (อนุมัติราคา) · WeeeU (รับทราบยอด)",
     "B2.2 เสนอราคา\ngold-lock 24h",
     "gold-lock 24h mock (visual only, logic = BE); escrow WeeeR→WeeeU; TODO BE: POST /scrap-offer",
     f"{BASE}/jobs/REP-001/scrap-offer"),

    # T-07 ─────────────────────────────────────────────────────────────────
    ("T-07", "ยืนยันเสร็จงาน (Complete Job)",
     "T-11 รายละเอียดงาน /jobs/[id]",
     "repairApi.complete(id) confirmed = true\nต้องกดปุ่มยืนยัน 1 ครั้ง",
     "T-01 /jobs (router.replace → job list)",
     "WeeeT (ช่าง) · WeeeR (รับทราบงานเสร็จ)",
     "ยืนยันปิดงาน\nconfirmed toggle",
     "repairApi.complete(); confirmed checkbox; router.replace('/jobs')",
     f"{BASE}/jobs/REP-001/complete"),

    # T-08 ─────────────────────────────────────────────────────────────────
    ("T-08", "รายงานตรวจสอบ (Pre-inspection)",
     "T-11 รายละเอียดงาน /jobs/[id]",
     "รูป ≥3 ≤5 · คลิป 0-3 (ไม่บังคับ) · notes\nไฟล์ ≤3MB (photo) / ≤30MB (video)",
     "T-11 /jobs/[id] (ส่งรายงานสำเร็จ)",
     "WeeeT (ช่าง) · WeeeR (รับรูปก่อนซ่อม) · WeeeU (รับทราบ)",
     "ถ่ายรูปก่อนซ่อม\nวิดีโอเพิ่มเติม\nหมายเหตุ",
     "repairApi.preInspect(); pre_inspection_photos + clips FormData; capture=environment",
     f"{BASE}/jobs/REP-001/inspect"),

    # T-09 ─────────────────────────────────────────────────────────────────
    ("T-09", "รายงานปัญหาเพิ่ม (M4 Issue)",
     "T-39 inspect (maintain) · T-37 progress",
     "MAINTAIN-M4-ISSUE state: draft / submitted\nรูปหลักฐาน ≤5 ใบ · cost breakdown อะไหล่+ค่าแรง",
     "submitted → รอลูกค้า (U-14 extra-cost)\ncancel → กลับ",
     "WeeeT (ช่าง) · WeeeU (U-14 extra-cost decision) · WeeeR (รับทราบ)",
     "M4 รายงานปัญหา\nBreakdown อะไหล่+ค่าแรง\nส่งรอลูกค้าอนุมัติ",
     "Mockup-only state-based; MAX_PHOTOS=5; ≠M5; TODO BE: POST /maintain/[id]/issue",
     f"{BASE}/jobs/MAINT-001/issue"),

    # T-10 ─────────────────────────────────────────────────────────────────
    ("T-10", "ของไม่ตรงประกาศ (S8 Mismatch)",
     "T-23 รับซาก /scrap/[id] (mismatch_reported)",
     "S8 state: pending (รอตัดสิน) / repriced (ปรับราคาแล้ว) / rejected (ปฏิเสธรับซาก)\nรูปหลักฐาน ≤5 ใบ",
     "pending→repriced (เสนอราคาใหม่ Gold Point)\nrejected→T-22 กลับ",
     "WeeeT (ช่าง) · WeeeR (อนุมัติ S8) · WeeeU (รับทราบผล)",
     "S8 ของไม่ตรง\nปรับราคาใหม่\nปฏิเสธ",
     "Mockup-only; state tabs 3 states; MAX_PHOTOS=5; TODO BE: POST /scrap/[id]/mismatch",
     f"{BASE}/jobs/SPJ-001/mismatch"),

    # T-11 ─────────────────────────────────────────────────────────────────
    ("T-11", "รายละเอียดงาน (Job Detail)",
     "T-01 /jobs · T-19 /today · T-18 Dashboard",
     "job status 12 states; service_type: on_site/pickup/parcel\naction buttons เปลี่ยนตาม status",
     "T-02 /diagnose · T-03 /repair · T-04 /pickup\nT-05 /schedule · T-06 /scrap-offer · T-07 /complete\nT-08 /inspect · T-33 /arrive · T-34 /depart\nT-35 /photo · T-36 /parts · T-37 /progress · T-38 /post-repair",
     "WeeeT (ช่าง)",
     "ดูรายละเอียด\naction buttons ตาม status",
     "repairApi.get(id); STATUS_LABEL 12 states; dynamic buttons; mock customer info",
     f"{BASE}/jobs/REP-001"),

    # T-12 ─────────────────────────────────────────────────────────────────
    ("T-12", "โปรไฟล์ช่าง",
     "Bottom Nav / T-18 Dashboard",
     "ช่างล็อกอินแล้ว\nแก้ไข: ชื่อ / specialties / education / postal code\nอัปโหลด cert ≤3MB ≤10 ไฟล์",
     "T-21 /settings\nT-18 Dashboard",
     "WeeeT (ช่าง)",
     "ดูโปรไฟล์\nแก้ไขข้อมูล\nอัปโหลด cert",
     "useAuth(); availableSpecialties; educationLevels; postalCodeMap; PaymentWalletCard; MAX_CERT=10 × 3MB",
     f"{BASE}/profile"),

    # T-13 ─────────────────────────────────────────────────────────────────
    ("T-13", "การแจ้งเตือน (Notifications)",
     "Bottom Nav / Header bell icon",
     "category: repair / schedule / point / system\nDARK theme (#1696F9 weeet-primary)",
     "แตะ notif → T-11 หรือ T-37 แล้วแต่ประเภท",
     "WeeeT (ช่าง)",
     "repair notif\nschedule notif\npoint notif\nsystem notif",
     "Mockup-only; 4 categories; read/unread state; DARK theme; pattern จาก weeeu/notifications",
     f"{BASE}/notifications"),

    # T-14 ─────────────────────────────────────────────────────────────────
    ("T-14", "พอยต์ของฉัน (Wallet)",
     "Bottom Nav / T-12 Profile",
     "Gold Point (ค่าธรรมเนียม/บริการ) · Silver Point (read-only)\nD91: ช่าง ไม่มี withdraw flow",
     "—",
     "WeeeT (ช่าง)",
     "ดู Gold balance\nดู Silver balance\nประวัติธุรกรรม",
     "D91 WeeeT=เติมเท่านั้น (ถอนไม่ได้); Silver read-only; DARK theme #1696F9; pattern จาก weeeu/wallet",
     f"{BASE}/wallet"),

    # T-15 ─────────────────────────────────────────────────────────────────
    ("T-15", "บันทึกผลซ่อมสำเร็จ (Repair Success)",
     "T-03 /jobs/[id]/repair (OTP confirm)",
     "Phase 3 Success Page\nOTP verified → repair complete",
     "T-01 /jobs (router.replace)",
     "WeeeT (ช่าง) · WeeeR (รับทราบซ่อมเสร็จ) · WeeeU (รับทราบ)",
     "ซ่อมเสร็จ success\nTier-1 A8 flow→success",
     "T-15 ScreenBadge; success state display; router.replace('/jobs')",
     f"{BASE}/jobs/REP-001/repair/success"),

    # ══════════════════════════════════════════════════════════════════════════
    # T-16..T-48 — Service Listings, Dashboard, Scrap, Parts, Maintain, Auth
    # ══════════════════════════════════════════════════════════════════════════

    # T-16 ─────────────────────────────────────────────────────────────────
    ("T-16", "รายละเอียดประกาศบริการ",
     "T-17 รายการประกาศ",
     "state D83: published / has_offer / matched / completed / cancelled\nescrow phase: none / held / released / refunded",
     "matched→completed (ยืนยันส่งมอบ ปล่อย escrow)\nmatched→cancelled (ยกเลิก คืนเงิน)\nT-17 กลับรายการ",
     "WeeeT (ช่าง)",
     "S-TA1 ส่งมอบ release escrow\nS-TA2 ยกเลิก refund",
     "Escrow Gold Point; D83 state machine; transitionListing()",
     f"{BASE}/listings/L-001"),

    # T-17 ─────────────────────────────────────────────────────────────────
    ("T-17", "ประกาศบริการของฉัน",
     "Bottom Nav / T-18 Dashboard",
     "GR-10 near-me filter (radius 20 km)\nfilter by state",
     "T-16 รายละเอียด",
     "WeeeT (ช่าง)",
     "ดูรายการ\nกรองใกล้ฉัน GR-10",
     "NearMeFilter (@app3r/ui) · roleTheme weeet-primary",
     f"{BASE}/listings"),

    # T-18 ─────────────────────────────────────────────────────────────────
    ("T-18", "Dashboard ช่าง (หน้าหลัก)",
     "Login success / Middleware redirect",
     "ช่างล็อกอินแล้ว",
     "T-01 /jobs · T-17 /listings · T-19 /today · T-12 /profile",
     "WeeeT (ช่าง)",
     "แสดงสถิติงานวันนี้",
     "สวัสดี + ชื่อช่าง + stats: งานวันนี้ / เสร็จ / งานรอ · mock data",
     f"{BASE}/dashboard"),

    # T-19 ─────────────────────────────────────────────────────────────────
    ("T-19", "งานวันนี้",
     "Bottom Nav / T-18 Dashboard",
     "jobs ของ tech-001 วันนี้ (scheduledAt today)\ncurrentStage = in_progress",
     "T-37 /jobs/[id]/progress",
     "WeeeT (ช่าง)",
     "active jobs\ndone jobs",
     "localStorage seed data; DailyQueueCard component; Phase D WebSocket (future)",
     f"{BASE}/today"),

    # T-20 ─────────────────────────────────────────────────────────────────
    ("T-20", "รายงาน / สรุปผลงาน",
     "Bottom Nav",
     "period: วันนี้ / สัปดาห์นี้ / เดือนนี้",
     "—",
     "WeeeT (ช่าง)",
     "สรุปงาน: total / completed / in_progress / assigned",
     "mock data; completion rate bar; recent jobs list; period selector",
     f"{BASE}/reports"),

    # T-21 ─────────────────────────────────────────────────────────────────
    ("T-21", "ตั้งค่า",
     "Bottom Nav / T-12 Profile",
     "ช่างล็อกอินแล้ว",
     "T-47 Logout → /login\nT-48 change-password",
     "WeeeT (ช่าง)",
     "GPS toggle\nnotification toggle\nlanguage th/en\nlogout\nเปลี่ยนรหัสผ่าน",
     "LocationPickerSection; PushPermissionBanner; changePassword()",
     f"{BASE}/settings"),

    # T-22 ─────────────────────────────────────────────────────────────────
    ("T-22", "รายการงานรับซาก",
     "Bottom Nav",
     "filter: ทั้งหมด / รับงานใหม่ / ถึงแล้ว / no-show / เสร็จสิ้น",
     "T-23 รายละเอียดงาน",
     "WeeeT (ช่างรับซาก)",
     "no-show alert badge (S9)\nactive count badge",
     "S9 no-show banner; S12 sourceRepairJobId badge · mock 3 jobs",
     f"{BASE}/scrap"),

    # T-23 ─────────────────────────────────────────────────────────────────
    ("T-23", "รับซาก — ตรวจสอบ + ดำเนินการ",
     "T-22 รายการงานรับซาก",
     "status: assigned / traveling / arrived / verifying / mismatch_reported / pickup_confirmed / no_show / completed / cancelled",
     "✅ ตรงตาม → pickup_confirmed\n⚠️ ไม่ตรง (S8) → mismatch_reported\n🚫 ไม่พบ (S9) → no_show\nT-22 กลับรายการ",
     "WeeeT (ช่างรับซาก)",
     "S8 ซากไม่ตรงสเปก\nS9 ไม่พบลูกค้า/ของ\nS12 repair badge",
     "GPS check-in; หลักฐานรูป S8; ราคาเสนอใหม่ Gold Point; no-show modal",
     f"{BASE}/scrap/SPJ-001"),

    # T-24 ─────────────────────────────────────────────────────────────────
    ("T-24", "อะไหล่ / วัสดุ (คลังร้าน)",
     "Bottom Nav",
     "search ชื่อ/SKU\nfilter category",
     "T-25 รายละเอียดอะไหล่\nT-30 ออเดอร์ของฉัน",
     "WeeeT (ช่าง)",
     "ดูคลัง\nค้นหาอะไหล่",
     "partsApi.list(); stock qty; condition: new/used/refurbished; reservedQty",
     f"{BASE}/parts"),

    # T-25 ─────────────────────────────────────────────────────────────────
    ("T-25", "รายละเอียดอะไหล่",
     "T-24 /parts",
     "part ID param",
     "T-28 cart\nT-36 /jobs/[id]/parts",
     "WeeeT (ช่าง)",
     "ดูรายละเอียด\nเพิ่มตะกร้า",
     "B5 อะไหล่; partsApi.get(id)",
     f"{BASE}/parts/p001"),

    # T-26 ─────────────────────────────────────────────────────────────────
    ("T-26", "แคตตาล็อกอะไหล่",
     "T-24 /parts",
     "—",
     "T-27 รายละเอียดในแคตตาล็อก",
     "WeeeT (ช่าง)",
     "เรียกดูแคตตาล็อก",
     "B5 รายการสั่งซื้อจากผู้ขาย",
     f"{BASE}/parts/catalog"),

    # T-27 ─────────────────────────────────────────────────────────────────
    ("T-27", "รายละเอียดในแคตตาล็อก",
     "T-26 /parts/catalog",
     "catalog item ID",
     "T-28 cart (เพิ่มตะกร้า)",
     "WeeeT (ช่าง)",
     "ดูรายละเอียดสินค้า\nเพิ่มตะกร้า",
     "B5 อะไหล่; catalog item detail",
     f"{BASE}/parts/catalog/CAT-001"),

    # T-28 ─────────────────────────────────────────────────────────────────
    ("T-28", "ตะกร้าอะไหล่",
     "T-27 catalog/[id]\nT-25 parts/[id]",
     "items in cart",
     "T-29 checkout",
     "WeeeT (ช่าง)",
     "ดูตะกร้า\nแก้จำนวน\nลบ",
     "localStorage cart; B5",
     f"{BASE}/parts/cart"),

    # T-29 ─────────────────────────────────────────────────────────────────
    ("T-29", "ชำระเงินสั่งซื้ออะไหล่",
     "T-28 cart",
     "cart items → order",
     "T-30 orders (order confirmed)",
     "WeeeT (ช่าง)",
     "ยืนยันสั่งซื้อ",
     "B5; สร้าง order; localStorage order IDs",
     f"{BASE}/parts/checkout"),

    # T-30 ─────────────────────────────────────────────────────────────────
    ("T-30", "ออเดอร์อะไหล่ของฉัน",
     "T-24 /parts\nT-29 checkout",
     "localStorage order IDs",
     "T-31 รายละเอียดออเดอร์",
     "WeeeT (ช่าง)",
     "ดูรายการออเดอร์",
     "B5; weeet_part_order_ids localStorage",
     f"{BASE}/parts/orders"),

    # T-31 ─────────────────────────────────────────────────────────────────
    ("T-31", "รายละเอียดออเดอร์อะไหล่",
     "T-30 /parts/orders",
     "order ID param",
     "T-30 กลับรายการ",
     "WeeeT (ช่าง)",
     "ดูสถานะ\nดูรายละเอียด",
     "B5 order detail; partsApi.getOrder()",
     f"{BASE}/parts/orders/ORD-001"),

    # T-32 ─────────────────────────────────────────────────────────────────
    ("T-32", "คำขออะไหล่ (Purchase Requests)",
     "T-24 /parts\nT-36 /jobs/[id]/parts",
     "purchase_request status",
     "—",
     "WeeeT (ช่าง)",
     "ดูรายการ purchase request",
     "B5; TODO BE: GET /api/v1/parts/purchase-requests/",
     f"{BASE}/parts/requests"),

    # T-33 ─────────────────────────────────────────────────────────────────
    ("T-33", "บันทึกถึงที่ (ซ่อม)",
     "T-11 job detail\nT-34 ออกเดินทาง",
     "ถ่ายรูป ≥2 (จุดสังเกต + กับลูกค้า)\nสูงสุด 3 ใบ · ≤3MB/ใบ",
     "T-11 /jobs/[id] (success)",
     "WeeeT (ช่าง)",
     "บังคับรูป 2 ใบ",
     "repairApi.arrive(); arrival_photos FormData; capture=environment",
     f"{BASE}/jobs/REP-001/arrive"),

    # T-34 ─────────────────────────────────────────────────────────────────
    ("T-34", "ออกเดินทาง (ซ่อม)",
     "T-11 job detail",
     "ต้องระบุ GPS (Geolocation API)",
     "T-11 /jobs/[id]",
     "WeeeT (ช่าง)",
     "GPS lat/lng required",
     "repairApi.depart(); departure_location {lat, lng}",
     f"{BASE}/jobs/REP-001/depart"),

    # T-35 ─────────────────────────────────────────────────────────────────
    ("T-35", "ถ่ายรูป / วิดีโอ",
     "T-37 /progress\nT-11 job detail",
     "ก่อนซ่อม / หลังซ่อม toggle",
     "กลับ (router.back)",
     "WeeeT (ช่าง)",
     "before toggle\nafter toggle\ncaption",
     "FileReader; ไม่ call API (mockup storage); capture=environment",
     f"{BASE}/jobs/REP-001/photo"),

    # T-36 ─────────────────────────────────────────────────────────────────
    ("T-36", "ขออะไหล่ (B5 สำรอง/สั่งซื้อ)",
     "T-37 /progress\nT-11 job detail",
     "source_type=warehouse: stock>0→สำรอง, stock=0→สั่งซื้อ\nsource_type=market: สั่งซื้อเสมอ",
     "submitted → submitted_reserve หรือ submitted_purchase (success state)",
     "WeeeT (ช่าง)",
     "B5 reserve (warehouse)\nB5 purchase request (market/stock=0)",
     "mockup no API; TODO BE: POST /reserve/ + /purchase-request/",
     f"{BASE}/jobs/REP-001/parts"),

    # T-37 ─────────────────────────────────────────────────────────────────
    ("T-37", "ติดตามงาน / อัปเดต",
     "T-19 /today\nT-11 job detail\nT-18 Dashboard",
     "ServiceProgress localStorage; currentStage; sub-stages",
     "T-35 /photo\nT-36 /parts\nT-11 กลับ",
     "WeeeT (ช่าง)",
     "C7 หยุดงานชั่วคราว\nC9 Dispute evidence\nprogress timeline\nstage update",
     "ServiceProgressTimeline; StepUpdateWizard; ProgressUpdateForm; evidencePhotos+GPS",
     f"{BASE}/jobs/REP-001/progress"),

    # T-38 ─────────────────────────────────────────────────────────────────
    ("T-38", "บันทึกหลังซ่อม",
     "T-37 /progress\nT-11 job detail",
     "รูปหลังซ่อม ≥3 · คลิป ≥1 (บังคับ)\nparts_used list",
     "T-11 /jobs/[id] (success)",
     "WeeeT (ช่าง)",
     "รูป before/after\nparts_used list",
     "repairApi.postRepair(); MIN_CLIPS=1; complete_photos+clips FormData",
     f"{BASE}/jobs/REP-001/post-repair"),

    # T-39 ─────────────────────────────────────────────────────────────────
    ("T-39", "ตรวจสภาพก่อนล้าง (M3)",
     "T-40 arrive (maintain)",
     "mode: normal / risk_form / risk_submitted / noshow_form / noshow_submitted",
     "normal → T-41 checklist\nrisk → WeeeR+WeeeU joint decision\nnoshow → T-46 / settled",
     "WeeeT (ช่างล้างเครื่อง)",
     "M3 ตรวจ\nM7 ไม่มาตามนัด\nD-Maintain-1 risk\nD-Maintain-2 convert-to-repair",
     "TODO BE C-4.1b: POST risk/ + convert-to-repair/ + no-show/",
     f"{BASE}/maintain/MAINT-001/inspect"),

    # T-40 ─────────────────────────────────────────────────────────────────
    ("T-40", "บันทึกถึงที่ ล้างเครื่อง (M2)",
     "T-43 depart (maintain)",
     "ถ่ายรูป ≥2 ก่อนล้าง · ≤3MB/ใบ",
     "T-39 inspect",
     "WeeeT (ช่างล้างเครื่อง)",
     "arrival_photos ≥2",
     "maintainApi.arrive(); arrival_photos FormData; เห็นสภาพก่อนล้าง",
     f"{BASE}/maintain/MAINT-001/arrive"),

    # T-41 ─────────────────────────────────────────────────────────────────
    ("T-41", "ล้างเครื่อง — Checklist (M4)",
     "T-39 inspect",
     "?type=general|deep|sanitize (searchParams)\nMaintain_CHECKLIST_ITEMS per type",
     "T-42 complete\nM9 terminate → WeeeR",
     "WeeeT (ช่างล้างเครื่อง)",
     "M4 ล้างตาม type\nM9 terminate (WeeeU ยุติงาน)",
     "MAINTAIN_CHECKLIST_ITEMS; partsApi.list() modal; terminate modal; maintainApi.complete(id)",
     f"{BASE}/maintain/MAINT-001/checklist"),

    # T-42 ─────────────────────────────────────────────────────────────────
    ("T-42", "เสร็จงาน ล้างเครื่อง (M5)",
     "T-41 checklist",
     "รูปหลังล้าง ≥2 + ลายเซ็นลูกค้า (บังคับ)",
     "T-01 /jobs (success · router.replace)",
     "WeeeT (ช่างล้างเครื่อง)",
     "M5 เสร็จ + sign\nSignaturePad",
     "SignaturePad (@/components/SignaturePad); customer_signature; maintainApi.complete()",
     f"{BASE}/maintain/MAINT-001/complete"),

    # T-43 ─────────────────────────────────────────────────────────────────
    ("T-43", "ออกเดินทาง ล้างเครื่อง (M1)",
     "T-01 /jobs หรือ Maintain job list",
     "GPS required (Geolocation API)",
     "T-40 arrive (maintain)\nT-01 กลับ",
     "WeeeT (ช่างล้างเครื่อง)",
     "M1 GPS departure_location",
     "maintainApi.depart(); {lat, lng}; background: weeet-primary/10",
     f"{BASE}/maintain/MAINT-001/depart"),

    # T-44 ─────────────────────────────────────────────────────────────────
    ("T-44", "ส่งคืนเครื่อง (Delivery flow)",
     "T-11 job detail (repair → delivery)",
     "ป้าย: en-route / completed / receipt\n(T-44 = ฐาน; ไม่มี base route ตรงๆ)",
     "completed → success\nreceipt → ลายเซ็นลูกค้า\nT-11 กลับ",
     "WeeeT (ช่าง)",
     "delivery en-route\ndelivery completed\ndelivery receipt",
     "3 sub-routes: /delivery/en-route · /delivery/completed · /delivery/receipt",
     f"{BASE}/jobs/REP-001/delivery/en-route"),

    # T-45 ─────────────────────────────────────────────────────────────────
    ("T-45", "ส่งซ่อมแบบพัสดุ (Parcel flow)",
     "T-11 job detail (scrap/parcel path)",
     "ป้าย: in-progress / tested\n(T-45 = ฐาน; ไม่มี base route ตรงๆ)",
     "tested → T-11 กลับ\nในแต่ละ sub-step ส่ง update",
     "WeeeT (ช่าง)",
     "parcel in-progress\nparcel tested",
     "2 sub-routes: /parcel/in-progress · /parcel/tested",
     f"{BASE}/jobs/REP-001/parcel/in-progress"),

    # T-46 ─────────────────────────────────────────────────────────────────
    ("T-46", "ถึงหน้างาน — ลูกค้าไม่อยู่/ไม่มาตามนัด (M7)",
     "T-39 inspect (mode = noshow_form)",
     "dev mockup artifact\npath มี /mockup segment → Phase 4 normalize ค่อยเปลี่ยน",
     "settle → T-01 /jobs",
     "WeeeT (ช่างล้างเครื่อง)",
     "M7 No-show\nลูกค้าไม่อยู่",
     "ScreenBadge T-46 label=ไม่มาตามนัด; path: /maintain/[id]/mockup/m7-noshow",
     f"{BASE}/maintain/MAINT-001/mockup/m7-noshow"),

    # T-47 ─────────────────────────────────────────────────────────────────
    ("T-47", "เข้าสู่ระบบ (Login)",
     "Middleware redirect (unauthenticated)\nLogout จาก T-21",
     "ยังไม่ล็อกอิน หรือ session หมดอายุ",
     "T-18 Dashboard (success)\nT-48 change-password (first-login)",
     "WeeeT (ช่าง)",
     "login form\nJWT auth\nfirst-login redirect",
     "นอก group (app); Auth flow; useAuth; middleware.ts",
     f"{BASE}/login"),

    # T-48 ─────────────────────────────────────────────────────────────────
    ("T-48", "เปลี่ยนรหัสผ่านครั้งแรก",
     "T-47 Login (must_change_password = true)",
     "first-login flag จาก backend",
     "T-18 Dashboard (success)",
     "WeeeT (ช่าง)",
     "first login\nforced password change",
     "Auth flow; Middleware check; changePassword(); นอก group (app)",
     f"{BASE}/change-password-first"),
]

assert len(SCREENS) == 48, f"Expected 48 screens, got {len(SCREENS)}"

# ── Build document ──────────────────────────────────────────────────────────
doc = Document()

# Page setup: landscape A4
section = doc.sections[0]
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(1.0)
section.right_margin = Cm(1.0)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("WeeeT — รายการหน้าจอ T-01 ถึง T-48 (48 จอ)")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x16, 0x96, 0xF9)

doc.add_paragraph(
    "แอพ WeeeT (ช่าง) · Port 3003 · สีหลัก #1696F9 · "
    "EXC-1 expanded T-01..T-48 · base origin/main da0e3a9"
).runs[0].font.size = Pt(9)

doc.add_paragraph()  # spacer

# Table
tbl = doc.add_table(rows=1, cols=9)
tbl.style = "Table Grid"

# Column widths (cm, total ~27.7 cm in landscape with margins)
COL_WIDTHS = [1.4, 4.0, 3.5, 4.5, 4.0, 2.5, 3.2, 3.5, 3.0]

WEEET_BLUE = RGBColor(0x16, 0x96, 0xF9)

# Header row
hdr_row = tbl.rows[0]
for i, (cell, hdr, w) in enumerate(zip(hdr_row.cells, HEADERS, COL_WIDTHS)):
    cell.width = Cm(w)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hdr)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Blue background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1696F9")
    tcPr.append(shd)

# Data rows
for row_idx, row_data in enumerate(SCREENS):
    row = tbl.add_row()
    fill = "E8F4FE" if row_idx % 2 == 0 else "FFFFFF"
    for i, (cell, val, w) in enumerate(zip(row.cells, row_data, COL_WIDTHS)):
        cell.width = Cm(w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i > 0 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(8)
        # Bold screen ID column
        if i == 0:
            run.bold = True
            run.font.color.rgb = WEEET_BLUE
        # Alternate row shading
        if fill != "FFFFFF":
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

out_path = r"D:\ClaudeCode\App3R\wt-p2p3-weeet\WeeeT_Screens_T01-T48.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
print(f"Rows: {len(tbl.rows)} (1 header + {len(tbl.rows)-1} data)")
